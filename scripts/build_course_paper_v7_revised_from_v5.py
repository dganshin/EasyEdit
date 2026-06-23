#!/usr/bin/env python3
"""Build the corrected v7 paper from the v5 academic draft plus frozen results.

This is intentionally different from build_course_paper_v7_final.py: it keeps
the richer v5 paper body and only replaces stale result/claim sections.
"""

from __future__ import annotations

import csv
import re
from pathlib import Path
from typing import Any, Dict, List, Sequence

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


BASE_MD = Path("docs/course_paper_v5_academic_polished.md")
TEMPLATE = Path("课程论文撰写模板 (研究类).docx")
OUT_MD = Path("docs/course_paper_v7_final_academic_polished.md")
OUT_DOCX = Path("课程论文_v7_final_academic_polished.docx")
CHANGELOG = Path("docs/COURSE_PAPER_V7_FINAL_CHANGELOG.md")
CHECKLIST = Path("docs/PAPER_FINAL_QUALITY_CHECKLIST.md")
TABLE_DIR = Path("artifacts/final_paper_assets_20260623/tables")
FIG_DIR = Path("artifacts/final_paper_assets_20260623/figures")


def read_csv(path: Path) -> List[Dict[str, str]]:
    with path.open(encoding="utf-8", newline="") as fh:
        return list(csv.DictReader(fh))


def md_table(rows: List[Dict[str, Any]], fields: Sequence[str]) -> str:
    lines = ["|" + "|".join(fields) + "|", "|" + "|".join(["---"] * len(fields)) + "|"]
    for row in rows:
        lines.append("|" + "|".join(str(row.get(f, "")) for f in fields) + "|")
    return "\n".join(lines)


def replace_between(text: str, start: str, end: str, replacement: str) -> str:
    i = text.index(start)
    j = text.index(end, i)
    return text[:i] + replacement.rstrip() + "\n\n" + text[j:]


def replace_through_line(text: str, start: str, end_line: str, replacement: str) -> str:
    i = text.index(start)
    j = text.index(end_line, i)
    k = text.find("\n", j)
    if k == -1:
        k = len(text)
    return text[:i] + replacement.rstrip() + "\n\n" + text[k:].lstrip()


def compact_synthetic() -> List[Dict[str, str]]:
    labels = {
        "Merged pre-edit": "合并后泄露模型",
        "Prompt Refusal": "提示拒答",
        "PACE_LITE_B20_K0": "PACE-Lite K0",
        "CAPE_ANCHOR_B20_K1": "CAPE-Anchor K1",
        "CAPE_ANCHOR_B20_K2": "CAPE-Anchor K2",
    }
    out = []
    for r in read_csv(TABLE_DIR / "table1_synthetic_privacy_main.csv"):
        method = r["Method"]
        out.append(
            {
                "方法": labels.get(method, method),
                "状态": r["Status"],
                "隐私值包含率↓": r["Private leak ↓"],
                "PII格式率↓": r["PII regex ↓"],
                "敏感格式率↓": r["Sensitive pattern ↓"],
                "隐私拒答率": r["Private refusal"],
                "公开知识保持率↑": r["Public contains ↑"],
            }
        )
    return out


def compact_cape() -> List[Dict[str, str]]:
    labels = {
        "PACE_LITE_B20_K0": "PACE-Lite K0（无公开锚点）",
        "CAPE_ANCHOR_B20_K1": "CAPE-Anchor K1",
        "CAPE_ANCHOR_B20_K2": "CAPE-Anchor K2",
    }
    out = []
    for r in read_csv(TABLE_DIR / "table2_cape_anchor_ablation.csv"):
        out.append(
            {
                "配置": labels.get(r["Variant"], r["Variant"]),
                "隐私值包含率↓": r["Private leak ↓"],
                "PII格式率↓": r["PII regex ↓"],
                "隐私拒答率": r["Private refusal"],
                "公开知识保持率↑": r["Public contains ↑"],
                "解释": "无锚点受限再编辑对照" if r["Variant"].endswith("K0") else ("较均衡折中" if r["Variant"].endswith("K1") else "公开保持更高但隐私压制更弱"),
            }
        )
    return out


def compact_qwen() -> List[Dict[str, str]]:
    out = []
    for r in read_csv(TABLE_DIR / "table3_qwen_public_transfer.csv"):
        if r["Status"].startswith("failed"):
            continue
        out.append(
            {
                "数据集": "CounterFact" if r["Dataset"] == "counterfact" else "zsRE",
                "方法": r["Method"].replace("ROME_PACE_EDIT", "ROME+PACE-Edit").replace("ROME_CAPE_EDIT", "ROME+CAPE-Edit"),
                "样本数": r["Cases"],
                "可靠性↑": r["Reliability ↑"],
                "泛化性↑": r["Generalization ↑"],
            }
        )
    return out


def compact_gptj() -> List[Dict[str, str]]:
    out = []
    for r in read_csv(TABLE_DIR / "table4_gptj_boundary_check.csv"):
        out.append(
            {
                "数据集": "CounterFact" if r["Dataset"] == "counterfact" else "zsRE",
                "方法": r["Method"].replace("ROME_PACE_EDIT", "ROME+PACE-Edit").replace("ROME_CAPE_EDIT", "ROME+CAPE-Edit"),
                "可靠性↑": r["Reliability ↑"],
                "泛化性↑": r["Generalization ↑"],
                "逐样本rewrite": r["Per-case rewrite"] or "未统计",
                "解释": "基础编辑正常" if r["Interpretation"] == "baseline normal" else "wrapper塌缩",
            }
        )
    return out


def compact_failures() -> List[Dict[str, str]]:
    out = []
    for r in read_csv(TABLE_DIR / "table5_failure_and_resource_limits.csv"):
        out.append({"范围": r["Scope"], "方法/现象": r["Method"], "状态": r["Status"], "原因": r["Reason"], "论文处理": r["Policy"]})
    return out


def build_result_block() -> str:
    synthetic = compact_synthetic()
    cape = compact_cape()
    qwen = compact_qwen()
    gptj = compact_gptj()
    failures = compact_failures()
    return f"""## 4.3  基础编辑方法对比分析

表 3 给出 V2 合成隐私基准上的主结果。与前文验证一致，合并后泄露模型同时具有较高隐私泄露和较高公开知识保持：隐私值包含率为 0.9387，公开知识保持率为 0.9766。提示拒答的隐私值包含率仍为 0.9490，说明仅在提示层要求模型拒答并不能改变模型内部已经形成的隐私事实关联。

表 3  Synthetic privacy 主结果（↓越低越好，↑越高越好）

{md_table(synthetic, ["方法", "状态", "隐私值包含率↓", "PII格式率↓", "敏感格式率↓", "隐私拒答率", "公开知识保持率↑"])}

从基础编辑器看，MEMIT 单轮编辑公开知识保持率达到 0.8472，但隐私值包含率仍为 0.8140，说明其在当前拒答式隐私清洗目标下更接近高保持、弱压制端点。ROME 将隐私值包含率降至 0.5787，公开知识保持率为 0.5591，处于隐私压制与知识保持之间的中间区域。FT 将隐私值包含率、PII 格式匹配率和敏感格式输出率均压到 0，但公开知识保持率仅为 0.0040，说明它更像强覆盖端点，而不是可直接部署的高质量清洗方案。

这一组结果保留了 v5 版本中的核心判断：标准知识编辑器直接迁移到隐私清洗任务时，会呈现不同的隐私—效用端点。不同方法的价值不应被压缩为单一排名，而应放入隐私压制、公开知识保持和拒答副作用的共同坐标系中解释。

## 4.4  闭环再编辑与副作用分析

PACE 和 CAPE 进一步证明，残余泄露驱动再编辑可以显著压低目标隐私输出。PACE 的隐私值包含率为 0.0243，CAPE 为 0.0443，均明显低于 ROME 和 MEMIT。然而，二者的公开知识保持率分别只有 0.0984 和 0.1119，说明强隐私压制伴随明显 public collapse。该现象是本文方法设计中的关键观察：闭环再编辑提高了攻击覆盖，但也可能把拒答目标扩散到同主体公开知识。

图 3 将各方法放入隐私压制—知识保持二维空间。横轴为 Public Contains，纵轴为 1 - Private Value Contains。理想方法应接近右上角，即公开知识保持率高且隐私值包含率低。当前没有方法同时占据该区域；FT、PACE 和 CAPE 位于强压制但低保持端，MEMIT 位于高保持但弱压制端，ROME 处于中间区域。

图 3  Synthetic privacy 隐私压制—知识保持权衡散点图

![图 3  Synthetic privacy 隐私压制—知识保持权衡散点图]({FIG_DIR / "fig3_synthetic_privacy_utility_tradeoff.png"})

## 4.5  CAPE-Anchor 消融分析

CAPE-Anchor 的作用不是继续追求最低泄露率，而是在闭环再编辑中显式加入公开知识锚点，检验能否把 PACE/CAPE 的公开知识塌缩拉回。表 4 给出有限消融结果。K0 是不加入 public anchor 的受限再编辑对照；K1 和 K2 逐步加入公开锚点。

表 4  CAPE-Anchor 消融结果

{md_table(cape, ["配置", "隐私值包含率↓", "PII格式率↓", "隐私拒答率", "公开知识保持率↑", "解释"])}

结果显示，PACE-Lite K0 的公开知识保持率为 0.4210，已经高于 naive PACE/CAPE；CAPE-Anchor K1 进一步提升到 0.6008，K2 提升到 0.6833。与此同时，隐私值包含率也从 0.3323 上升到 0.4603 和 0.6357。该变化说明公开锚点确实改变了闭环编辑的操作点，但它不是“免费”的改进：公开知识保持越强，隐私压制强度越弱。

因此，本文最终采用限定的 Claim A：CAPE-Anchor 形成有限有效改进。其含义不是 CAPE-Anchor 隐私泄露最低，也不是完全解决隐私清洗；而是相对于 naive PACE/CAPE，它在保留可观隐私压制能力的同时显著提高公开知识保持，使方法从强拒答端点移动到更可解释的 privacy-utility trade-off 区域。

图 4  CAPE-Anchor 消融折中图

![图 4  CAPE-Anchor 消融折中图]({FIG_DIR / "fig4_cape_anchor_ablation_tradeoff.png"})

## 4.6  分攻击类型分析

分攻击类型分析仍保留 v5 版本中的观察：补全诱导通常比直接询问更难清洗，因为续写式上下文不一定触发与问答式请求相同的拒答模式。该现象说明，隐私清洗不能只面向 canonical question 构造编辑请求，还需要覆盖 paraphrase、completion、roleplay 和 context 等不同攻击路径。本文的 PACE/CAPE 设计正是从残余泄露样本回流到请求构造阶段，用于补足单轮 direct-only 请求的覆盖不足。

从安全解释看，敏感格式输出与目标值泄露也需要区分。模型可能没有输出目标电话号码，却生成另一个电话号码格式字符串；这种行为不构成目标值精确泄露，但在隐私安全场景中仍然提示模型保持敏感格式生成倾向。因此，本文同时报告隐私值包含率、PII 格式匹配率和敏感格式输出率，而不是只报告单一拒答率。

## 4.7  Qwen 公开事实编辑迁移验证

CounterFact 和 zsRE 不是 PII 清洗任务。本文将其作为公开事实编辑迁移验证，用于观察 closed-loop request selection 在非隐私 factual editing 场景下的外部行为。表 5 显示，Qwen 上 ROME/FT 是较强 baseline，wrapper 可以迁移运行但效果中等。CounterFact 上 ROME+PACE-Edit 和 ROME+CAPE-Edit 的可靠性均为 0.5179；zsRE 上二者可靠性均为 0.3117。

表 5  Qwen public transfer 结果

{md_table(qwen, ["数据集", "方法", "样本数", "可靠性↑", "泛化性↑"])}

这一结果不能写成“公开数据集证明隐私清洗成功”。更准确的解释是：PACE/CAPE 的闭环请求选择思想可以被抽象到公开 factual editing 任务中运行，但在公开基准上的收益并不强，且不优于强 ROME/FT baseline。因此，Qwen public transfer 在本文中属于外部行为验证和压力测试，而不是主贡献证据。

图 5  Qwen public transfer 对比图

![图 5  Qwen public transfer 对比图]({FIG_DIR / "fig5_public_transfer_qwen.png"})

## 4.8  GPT-J 第二模型边界验证

为检查公开迁移实验是否完全依赖 Qwen2.5-7B，本文在 GPT-J-6B 上补充第二模型公开事实编辑验证。结果显示，GPT-J-6B 上 ROME 与 FT baseline 均能正常运行，在 CounterFact 与 zsRE 上获得较高 rewrite 成功率。然而，将未重新调参的 ROME-based PACE/CAPE wrapper 直接迁移到 GPT-J-6B 时，wrapper 出现明显塌缩。

表 6  GPT-J second-model boundary check

{md_table(gptj, ["数据集", "方法", "可靠性↑", "泛化性↑", "逐样本rewrite", "解释"])}

per-case 审计进一步排除了路径混用和汇总脚本误读的可能：CounterFact wrapper 的 rewrite 成功数为 0/220，zsRE wrapper 为 1/216，而同一模型上的 ROME/FT baseline 正常。该结果不应作为第二模型有效性的正向证据，也不能否定 Qwen synthetic privacy 主实验结论。它说明闭环请求扩展策略的跨模型迁移依赖底层编辑器超参数、请求集合规模和局部性约束，后续需要模型特定校准或 retain-aware objective。

图 6  GPT-J wrapper boundary check 图

![图 6  GPT-J wrapper boundary check 图]({FIG_DIR / "fig6_gptj_wrapper_boundary_check.png"})

## 4.9  失败项与资源限制分析

表 7 汇总 failed、resource-limited 与 negative boundary 项。KN 在部分设置中受 48GB GPU 显存限制，IKE 缺少 SentenceTransformer 依赖；GPT-J wrapper 虽然运行完成，但作为负向边界结果处理。这些内容不应隐藏，也不应被当作有效 0 值；它们体现的是当前资源条件和方法迁移边界。

表 7  Failure and resource limitation

{md_table(failures, ["范围", "方法/现象", "状态", "原因", "论文处理"])}

图 7  Failure/resource limitation summary

![图 7  Failure/resource limitation summary]({FIG_DIR / "fig7_failure_matrix_summary.png"})

## 4.10  有效性与局限性分析

本文实验的内部有效性来自统一模型、统一数据和统一评估协议。主要局限在于合成人物规模有限，公开事实类型仍较少，自动 contains 指标只能近似衡量目标值是否出现，尚需人工语义评估补充。真实预训练隐私泄露还涉及数据来源追溯、合规边界和跨语境攻击，需在更复杂数据上进一步验证。

外部有效性方面，Qwen public transfer 说明 wrapper 能够迁移运行但效果中等；GPT-J boundary check 进一步说明，未重新调参的 wrapper 并不能稳定跨模型迁移。换言之，本文的主结论应限定在 Qwen synthetic privacy 主实验中：CAPE-Anchor 提供更合理的隐私压制—知识保持折中，而不是宣称一个跨模型稳定成功的通用算法。

构念有效性方面，公开知识包含式保持率对回答质量的刻画仍然有限，公开问题拒答率也依赖拒答标记集合。后续可以引入语义匹配模型、人类评估和安全分类器，对回答正确性、拒答合理性和敏感性进行更细粒度标注。这样可以把本文的自动评估框架扩展为更完整的多层评估体系。"""


def build_markdown() -> str:
    text = BASE_MD.read_text(encoding="utf-8")
    text = replace_through_line(
        text,
        "# 摘要",
        "关键词：大语言模型；模型编辑；隐私知识清洗；知识保持；过度拒答；ROME；MEMIT；PACE；CAPE",
        """# 摘要

大语言模型在预训练和后训练过程中可能形成对个人敏感信息的参数化记忆，并在特定提示下输出电话号码、邮箱等个人可识别信息（Personally Identifiable Information，PII）。仅依赖输出过滤或提示层拒答难以改变模型内部知识关联，而全量重训成本较高。针对这一问题，本文研究基于模型编辑的隐私知识清洗方法，目标是在降低目标隐私泄露的同时尽量保持同主体公开知识、同关系其他主体公开知识和通用知识。本文首先构建 private/public 解耦的合成隐私评测基准，通过低秩适配（Low-Rank Adaptation，LoRA）向 Qwen2.5-7B 注入可控隐私记忆，得到统一的合并后隐私泄露模型。随后，基于 EasyEdit 框架比较秩一模型编辑（Rank-One Model Editing，ROME）、Transformer 批量记忆编辑（Mass-Editing Memory in a Transformer，MEMIT）、微调（Fine-Tuning，FT）以及隐私感知闭环编辑（Privacy-Aware Closed-loop Editing，PACE）的行为差异，并进一步设计副作用感知隐私编辑（Collateral-Aware Privacy Editing，CAPE）和引入公开知识锚点的 CAPE-Anchor。实验结果表明，提示拒答难以改变模型内部隐私记忆；MEMIT 保留公开知识较好但隐私压制不足，ROME 呈现中间权衡区域，FT、PACE 和 CAPE 构成强隐私压制但高副作用端点；CAPE-Anchor 通过显式加入公开知识锚点，将 naive PACE/CAPE 中接近塌缩的公开知识保持能力拉回，在牺牲部分隐私压制强度的情况下形成更合理的隐私压制—知识保持折中。公开事实编辑实验进一步显示，该闭环请求选择思想可以在 Qwen 上迁移运行但效果中等；GPT-J-6B 上的边界实验表明，未重新调参的 wrapper 对模型结构和编辑器超参数较敏感。本文结果为后续研究模型特定校准、局部性约束和 retain-aware objective 提供了可复现实验依据。

关键词：大语言模型；模型编辑；隐私知识清洗；知识保持；过度拒答；ROME；MEMIT；PACE；CAPE-Anchor""",
    )
    text = replace_between(
        text,
        "## 1.4  本文研究内容与贡献",
        "## 1.5  论文结构",
        """## 1.4  本文研究内容与贡献

本文围绕“基于模型编辑的大语言模型隐私知识清洗”构建完整实验流程。首先，构造包含隐私事实和公开事实的合成数据集；其次，通过 LoRA 注入形成统一的隐私泄露模型；再次，在同一模型和同一评估协议下比较 ROME、MEMIT、FT、PACE、CAPE 和 CAPE-Anchor；最后，从隐私值包含率、公开知识包含式保持率、公开问题拒答率以及公开事实编辑迁移验证等维度分析各方法的行为边界。

本文贡献包括四点。第一，构建 private/public 解耦的隐私清洗评测基准，使目标隐私压制和公开知识保持能够同时度量。第二，基于 LoRA 构造可控泄露模型，并在 EasyEdit 框架下将 ROME、MEMIT、FT 等编辑方法适配到隐私拒答式清洗任务。第三，提出 PACE、CAPE 与 CAPE-Anchor 三类 wrapper 层请求选择策略，用于分析残余泄露闭环再编辑、过度拒答和公开知识锚点约束。第四，建立隐私压制—知识保持权衡分析框架，将 Public Refusal、attack-type split、Qwen public transfer 和 GPT-J boundary check 纳入同一结果口径。

在方法定位上，PACE、CAPE 和 CAPE-Anchor 均属于编辑器外层策略，而非新的底层参数更新算法。这样的设计使实验能够保持 ROME/MEMIT 原始实现相对干净，同时把研究重点放在隐私任务特有的请求选择和副作用控制上。本文最终主张也相应保持克制：CAPE-Anchor 不是全面优于所有方法的新编辑器，而是在 Qwen synthetic privacy 主任务上形成了更合理的隐私压制—知识保持折中。""",
    )
    text = replace_between(text, "## 4.3  基础编辑方法对比分析", "# 5  系统设计与软件实现", build_result_block())
    text = text.replace(
        "![图 1 隐私知识清洗实验框架图](D:/BaiduSyncdisk/repo/EasyEdit/主要框架大图.png)",
        f"![图 1 隐私知识清洗实验框架图]({FIG_DIR / 'fig1_pipeline_overview.png'})",
    )
    text = text.replace(
        "![图 1  隐私知识清洗总体框架](D:/BaiduSyncdisk/repo/EasyEdit/主要框架大图.png)",
        f"![图 1  隐私知识清洗总体框架]({FIG_DIR / 'fig1_framework_pipeline.png'})",
    )
    text = text.replace(
        "![图 2 PACE 与 CAPE 请求选择流程图](D:/BaiduSyncdisk/repo/EasyEdit/CAPEv0.png)",
        f"![图 2 PACE 与 CAPE 请求选择流程图]({FIG_DIR / 'fig2_pace_cape_anchor_flow.png'})",
    )
    text = text.replace(
        "![图 2  CAPE 请求筛选机制示意图](D:/BaiduSyncdisk/repo/EasyEdit/CAPEv0.png)",
        f"![图 2  PACE/CAPE-Anchor 请求选择机制示意图]({FIG_DIR / 'fig2_pace_cape_anchor_mechanism.png'})",
    )
    text = text.replace("图 2  CAPE 请求筛选机制示意图", "图 2  PACE/CAPE-Anchor 请求选择机制示意图")
    text = text.replace("表 1  V2 合成隐私评测基准统计\n\n表 1  V2 合成隐私评测基准统计", "表 1  V2 合成隐私评测基准统计")
    text = re.sub(
        r"\n图 7  实验流水线图\n\n!\[图 7  实验流水线图\]\(D:/BaiduSyncdisk/repo/EasyEdit/docs/figures_v5/fig7_pipeline\.png\)\n\n图 8  指标体系图\n\n!\[图 8  指标体系图\]\(D:/BaiduSyncdisk/repo/EasyEdit/docs/figures_v5/fig8_metrics\.png\)\n",
        "\n",
        text,
    )
    text = text.replace(
        "每轮实验保存请求文件、case id、评估结果、摘要文件和运行日志。服务器负责 GPU 编辑与生成，本地负责脚本维护、结果分析和论文生成。",
        "每轮实验保存请求文件、case id、评估结果、摘要文件和运行日志。GPU 实验环境负责模型编辑与批量生成，代码管理环境负责脚本维护、结果分析和论文生成。",
    )
    text = text.replace(
        "实验流程采用本地生成脚本、服务器执行实验、本地汇总分析的协同方式。所有关键实验均保留请求文件、评估输出、摘要文件和运行日志。该设计保证不同编辑策略能够在同一数据和同一评估口径下比较。",
        "实验流程采用脚本化数据构建、GPU 环境批量执行和本地结果归档的协同方式。所有关键实验均保留请求文件、评估输出、摘要文件和运行日志。该设计保证不同编辑策略能够在同一数据和同一评估口径下比较。",
    )
    text = text.replace("## 5.3  本地与服务器协同机制", "## 5.3  实验环境与版本管理机制")
    text = text.replace(
        "本地 Windows 仓库用于代码和文档管理，AutoDL Linux 服务器用于运行大模型实验。同步方式采用 git pull/push，避免在服务器上进行大规模手改。",
        "代码管理环境用于脚本维护、文档生成和结果分析，Linux GPU 环境用于运行大模型编辑、批量生成和评估。同步方式采用 Git 版本控制，避免在实验环境中进行不可追踪的大规模手改。",
    )
    text = text.replace(
        "本地与服务器分工还降低了实验过程中的风险。本地环境适合进行脚本修改、论文生成和结果分析；服务器环境适合加载大模型、执行编辑和批量生成。通过 git 同步代码，通过 artifacts 同步小型结果，可以避免在服务器上临时修改导致版本不可追踪。",
        "这种分工降低了实验过程中的版本风险。代码管理环境适合进行脚本修改、论文生成和结果分析；GPU 环境适合加载大模型、执行编辑和批量生成。通过 Git 同步代码，通过 artifacts 归档结果，可以避免临时修改导致版本不可追踪。",
    )
    text = replace_between(
        text,
        "# 6  总结与展望",
        "# 参考文献",
        """# 6  总结与展望

本文围绕基于模型编辑的大语言模型隐私知识清洗问题，构建了从合成隐私评测基准、LoRA 可控隐私注入、模型编辑执行到副作用审计的完整实验闭环。该流程使待清洗隐私事实、需要保留的公开事实以及多类攻击提示均具有明确 ground truth，从而能够在统一协议下比较不同编辑策略。

实验表明，不同编辑策略在隐私压制和知识保持两个目标上呈现差异化端点。提示拒答基本不能改变模型内部隐私记忆；MEMIT 表现出较强公开知识保持能力，但在拒答式隐私清洗目标下压制强度有限；ROME 在隐私值输出压制和公开知识保持之间形成中间区域；FT、PACE 和 CAPE 构成强隐私压制但公开保持较弱的端点。

本文进一步引入公开问题拒答率作为副作用指标，用于区分合理隐私拒答和公开知识误伤。CAPE-Anchor 在闭环请求构造中加入公开知识锚点，使 naive PACE/CAPE 中接近塌缩的公开知识保持能力被拉回。最终结果支持限定的 Claim A：CAPE-Anchor 在 Qwen synthetic privacy 主任务上形成有限有效改进，即提供更合理的隐私压制—知识保持折中，但不是泄露最低的方法，也不是完全解决隐私清洗。

公开实验揭示了方法的迁移边界。Qwen CounterFact/zsRE 结果说明 ROME-based wrapper 可以迁移运行，但效果中等；GPT-J-6B 结果则显示，未重新调参的 wrapper 在第二模型上几乎塌缩，而 ROME/FT baseline 正常。该边界结果表明，闭环请求扩展策略对底层编辑器超参数、模型结构和请求集合规模敏感。

后续研究可以从三个方向推进。第一，引入模型特定校准和 locality-aware constraint，使闭环 wrapper 在不同模型上更稳定。第二，在编辑目标中加入 retain-aware objective，把公开知识锚点从请求构造扩展为显式优化约束。第三，结合人工评估和更多真实场景数据，对拒答合理性、公开回答质量和多轮攻击鲁棒性进行更细粒度分析。""",
    )
    text = text.replace("CAPE 配置消融对比图", "CAPE-Anchor 消融折中图")
    return text[text.index("# 摘要") :]


def clear_document(doc: Document) -> None:
    body = doc._body._element
    for child in list(body):
        if child.tag == qn("w:sectPr"):
            continue
        body.remove(child)


def set_run_font(run, font="宋体", size=10.5, bold=False, color=None) -> None:
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def add_paragraph(doc: Document, text: str, style: str | None = None, first_indent=True, align=None) -> None:
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    elif first_indent:
        p.paragraph_format.first_line_indent = Pt(21)
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_run_font(run)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_table(doc: Document, rows: List[List[str]]) -> None:
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row in enumerate(rows):
        for j, value in enumerate(row):
            cell = table.cell(i, j)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if len(value) < 16 else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(value)
            set_run_font(run, size=8.5, bold=(i == 0))
            if i == 0:
                set_cell_shading(cell, "E8EEF5")
    doc.add_paragraph()


def add_abstract_table(doc: Document, abstract: str, keywords: str) -> None:
    table = doc.add_table(rows=2, cols=1)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, value in enumerate([f"摘要：\n{abstract}", keywords]):
        cell = table.cell(i, 0)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.text = ""
        p = cell.paragraphs[0]
        p.paragraph_format.line_spacing = 1.25
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(value)
        set_run_font(run, size=10.5, bold=False)
    doc.add_paragraph()


def parse_md_table(lines: List[str], i: int) -> tuple[List[List[str]], int]:
    rows = []
    while i < len(lines) and lines[i].strip().startswith("|"):
        line = lines[i].strip()
        if not re.match(r"^\|[-:| ]+\|$", line):
            rows.append([c.strip() for c in line.strip("|").split("|")])
        i += 1
    return rows, i


def add_image(doc: Document, path: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(Path(path)), width=Cm(14.8))


def markdown_to_docx(md: str, out_path: Path) -> None:
    doc = Document(TEMPLATE)
    clear_document(doc)
    for sec in doc.sections:
        sec.top_margin = Cm(2.54)
        sec.bottom_margin = Cm(2.54)
        sec.left_margin = Cm(2.8)
        sec.right_margin = Cm(2.6)

    # Template-style cover.
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("课程论文")
    set_run_font(run, "黑体", 22, True)
    for text in [
        "课  程：智能系统综合实践",
        "题  目：基于模型编辑的大语言模型隐私知识清洗与保护研究",
        "级、专业：        级  计算机科学与技术  专业        班",
        "学生姓名：",
        "提交日期：2026 年 6 月 23 日",
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        set_run_font(run, "宋体", 14, False)
    doc.add_page_break()

    lines = md.splitlines()
    i = 0
    seen_heading = False
    abstract_done = False
    while i < len(lines):
        line = lines[i].rstrip()
        if not line:
            i += 1
            continue
        if not seen_heading and not line.startswith("#"):
            i += 1
            continue
        if line.startswith("# "):
            seen_heading = True
            title = line[2:].strip()
            if title == "课程论文":
                i += 1
                continue
            if title == "摘要" and not abstract_done:
                i += 1
                abstract_parts: List[str] = []
                keywords = ""
                while i < len(lines):
                    current = lines[i].strip()
                    if current.startswith("关键词："):
                        keywords = current
                        i += 1
                        break
                    if current:
                        abstract_parts.append(current)
                    i += 1
                add_abstract_table(doc, "\n".join(abstract_parts), keywords)
                abstract_done = True
                continue
            p = doc.add_heading(level=1)
            run = p.add_run(title)
            set_run_font(run, "黑体", 16, True, RGBColor(46, 116, 181))
            i += 1
            continue
        if line.startswith("## "):
            p = doc.add_heading(level=2)
            run = p.add_run(line[3:].strip())
            set_run_font(run, "黑体", 13, True, RGBColor(46, 116, 181))
            i += 1
            continue
        if line.startswith("!["):
            m = re.match(r"!\[(.*?)\]\((.*?)\)", line)
            if m:
                add_image(doc, m.group(2))
            i += 1
            continue
        if line.startswith("|"):
            rows, i = parse_md_table(lines, i)
            add_table(doc, rows)
            continue
        if line.startswith("$$"):
            add_paragraph(doc, line.strip("$"), first_indent=False, align=WD_ALIGN_PARAGRAPH.CENTER)
            i += 1
            continue
        add_paragraph(doc, line)
        i += 1
    doc.save(out_path)


def write_logs() -> None:
    CHANGELOG.write_text(
        """# Course Paper V7 Final Changelog

Last updated: 2026-06-23
Final result status: frozen
No further GPU expansion unless explicitly approved

## 本次修正

1. 退回 v5 学术论文版本作为正文母版，保留其绪论、相关工作、方法设计和系统实现的论文化表达。
2. 删除 v6/v7 中的结果占位、工程日志式描述和过期扩展计划。
3. 将第4章结果段替换为 frozen final artifacts：synthetic 主结果、CAPE-Anchor 消融、Qwen public transfer、GPT-J boundary check 和 failure/resource limitation。
4. DOCX 从 `课程论文撰写模板 (研究类).docx` 生成封面与正文样式，不覆盖旧版文档。
5. GPT-J 结果固定为边界证据，不写成跨模型成功。
""",
        encoding="utf-8",
    )
    CHECKLIST.write_text(
        """# Paper Final Quality Checklist

Last updated: 2026-06-23
Final result status: frozen
No further GPU expansion unless explicitly approved

## 内容检查

- [x] 以 v5 学术论文为主体，不是重新生成实验报告。
- [x] 主实验限定为 Qwen synthetic privacy。
- [x] CAPE-Anchor 写成有限有效改进。
- [x] Qwen public 写成公开事实编辑迁移验证。
- [x] GPT-J 写成 second-model boundary check。
- [x] failed/missing/resource-limited 明确列入表格。
- [x] 主论文未保留“待填/占位/服务器/Codex/运行脚本”等过程词。

## 格式检查

- [x] DOCX 使用 `课程论文撰写模板 (研究类).docx` 作为生成基底。
- [x] 包含模板式封面字段。
- [x] 包含 7 张表和最终结果图。
- [ ] 本机未检测到 LibreOffice/soffice，仍需用 Word 打开进行最终视觉检查。
- [ ] 学生姓名、年级专业、提交日期可按实际提交要求手工补齐。
""",
        encoding="utf-8",
    )


def main() -> int:
    md = build_markdown()
    OUT_MD.write_text(md, encoding="utf-8")
    markdown_to_docx(md, OUT_DOCX)
    write_logs()
    print(f"wrote {OUT_MD}")
    print(f"wrote {OUT_DOCX}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
