import csv
import json
import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE_DOCX = ROOT / "课程论文_v2.docx"
TEMPLATE_DOCX = ROOT / "课程论文撰写模板 (研究类).docx"
BACKUP_DOCX = ROOT / "课程论文_v2_backup_before_rewrite.docx"
OUT_DOCX = ROOT / "课程论文_v2_rewritten.docx"
DRAFT_MD = ROOT / "docs" / "course_paper_rewrite_draft.md"
LOG_MD = ROOT / "docs" / "COURSE_PAPER_REWRITE_LOG.md"

ANALYSIS = ROOT / "artifacts" / "analysis_v2_audit_20260622"
MEMIT_DIR = ROOT / "artifacts" / "run_20260622_v2_memit_direct"
CAPE_DIR = ROOT / "artifacts" / "run_20260622_v2_cape_b1_tau05"
TRADEOFF_PNG = ANALYSIS / "privacy_utility_tradeoff.png"


def read_csv(path: Path):
    with path.open("r", encoding="utf-8") as fh:
        return list(csv.DictReader(fh))


def read_json(path: Path):
    with path.open("r", encoding="utf-8") as fh:
        return json.load(fh)


def fmt(value, digits=4):
    return f"{float(value):.{digits}f}"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(str(text))
    run.bold = bold
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(9)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, rows, headers):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True)
        set_cell_shading(table.rows[0].cells[i], "E8EEF5")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
    doc.add_paragraph()
    return table


def set_doc_styles(doc):
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)

    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.space_after = Pt(6)

    for style_name, size in [("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 12)]:
        style = doc.styles[style_name]
        style.font.name = "黑体"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor(31, 77, 120)
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)


def add_para(doc, text, style=None, bold_prefix=None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.first_line_indent = Pt(21) if style is None else None
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        r1.bold = True
        r2 = p.add_run(text[len(bold_prefix):])
    else:
        p.add_run(text)
    for run in p.runs:
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        run.font.size = Pt(10.5)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def build_markdown(data):
    mc = data["method_comparison"]
    over = data["over_refusal"]
    cape = data["cape_report"]
    lines = [
        "# 基于模型编辑的大语言模型隐私知识清洗与保护研究",
        "",
        "## 摘要",
        "",
        "大语言模型在预训练与后训练过程中可能记忆并输出个人隐私信息。本文围绕 Privacy Knowledge Sanitization via Model Editing 构建研究型课程项目：首先设计包含 private/public 对照的 synthetic privacy benchmark；然后使用 MLP-only LoRA 向 Qwen2.5-7B 注入可控隐私记忆，形成 merged leakage model；进一步基于 EasyEdit 框架实现 ROME、MEMIT 与 PACE（Privacy-Aware Closed-loop Editing）等编辑策略，并从 private leakage、public retain、over-refusal 与 privacy-utility trade-off 四个维度进行审计。实验结果表明，现有编辑策略在隐私压制和公开知识保持之间存在明显权衡：ROME 能降低部分泄露但损伤公开知识，MEMIT 保留公开知识较好但隐私压制不足，PACE 能显著压低泄露但带来严重过度拒答。基于上述发现，本文进一步提出 CAPE-v0（Collateral-Aware Privacy Editing）作为副作用感知的请求选择方案，为后续降低 collateral damage 提供改进方向。",
        "",
        "关键词：大语言模型；模型编辑；隐私知识清洗；知识保持；过度拒答；ROME；MEMIT；PACE",
        "",
        "## 主结果表",
        "",
        "| 方法 | Private Value Contains | PII-format Regex | Sensitive Pattern | Private Refusal | Public Contains | Public Refusal |",
        "| --- | ---: | ---: | ---: | ---: | ---: | ---: |",
    ]
    over_map = {r["method_key"]: r for r in over}
    for r in mc:
        o = over_map[r["method_key"]]
        lines.append(
            f"| {r['method']} | {fmt(r['private_exact'])} | {fmt(r['private_regex'])} | {fmt(r['private_sensitive'])} | {fmt(r['private_refusal'])} | {fmt(r['public_contains'])} | {fmt(o['public_refusal_rate'])} |"
        )
    lines.extend([
        "",
        "## CAPE-v0 当前选择结果",
        "",
        f"CAPE-v0 在 `B=1, tau=0.5` 设置下，从 {cape['selection_summary']['num_candidates']} 个 residual leakage candidates 中选出 {cape['selection_summary']['num_selected']} 条 round2 re-edit requests。其中 public-anchor blocking 跳过 {cape['selection_summary']['skipped_public_anchor']} 条，per-person budget 跳过 {cape['selection_summary']['skipped_budget']} 条。该结果目前仅代表请求选择阶段，不代表 CAPE 模型编辑实验已经完成。",
    ])
    return "\n".join(lines) + "\n"


def add_cover(doc):
    for _ in range(3):
        doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("课程论文")
    run.bold = True
    run.font.name = "黑体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    run.font.size = Pt(26)
    for _ in range(3):
        doc.add_paragraph()
    for line in [
        "课  程：智能系统综合实践",
        "题  目：基于模型编辑的大语言模型隐私知识清洗与保护研究",
        "级、专业：          级          专业          班",
        "学生姓名：蒋兴",
        "提交日期：2026 年 6 月 22 日",
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(line)
        r.font.name = "宋体"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        r.font.size = Pt(14)
    doc.add_page_break()


def add_abstract(doc):
    heading(doc, "摘要", 1)
    add_para(
        doc,
        "大语言模型在预训练、指令微调和后训练过程中可能记忆并输出个人隐私信息。传统提示词约束、输出过滤和安全拒答模板主要工作在输入输出层，难以直接改变模型内部已经形成的事实记忆。本文围绕“基于模型编辑的大语言模型隐私知识清洗与保护”开展研究型课程设计，目标是在可控实验环境中分析模型编辑方法对目标隐私事实的压制能力及其对公开知识保持的副作用。项目首先构建包含 private/public 对照关系的 synthetic privacy benchmark，并通过 MLP-only LoRA 向 Qwen2.5-7B 注入可控隐私记忆，得到能够稳定泄露目标隐私值的 merged leakage model。在此基础上，本文使用 EasyEdit 框架实现 ROME 与 MEMIT direct-only 隐私拒答编辑，并设计 PACE（Privacy-Aware Closed-loop Editing，隐私感知闭环编辑）策略对 residual leakage 进行再编辑。",
    )
    add_para(
        doc,
        "实验评估覆盖 Private Value Contains、PII-format Regex、Sensitive Pattern、Private Refusal、Public Contains 和 Public Refusal 等指标。审计结果表明，原脚本中的 private_exact 实际为目标隐私值的 substring contains 匹配，不应解释为严格 exact equality；public_contains 也属于宽松包含式保持指标，不等价于严格事实正确率。在统一口径下，ROME 能将 Private Value Contains 从 0.9387 降至 0.5787，但 Public Contains 同时降至 0.5591；MEMIT 的 Public Contains 达到 0.8472，但 Private Value Contains 仍为 0.8140，隐私压制不足；PACE 能将泄露压到很低，但 Public Refusal 最高达到 0.9675，表明 naive residual re-edit 容易导致过度拒答。基于上述发现，本文进一步提出 CAPE-v0（Collateral-Aware Privacy Editing，副作用感知隐私编辑）作为后续改进方案，通过 public-anchor blocking 与 per-person edit budget 控制再编辑副作用。",
    )
    add_para(doc, "关键词：大语言模型；模型编辑；隐私知识清洗；知识保持；过度拒答；ROME；MEMIT；PACE")


def add_chapter_1(doc):
    heading(doc, "1  前言", 1)
    heading(doc, "1.1  研究背景和意义", 2)
    add_para(doc, "大语言模型（Large Language Models, LLMs）已在问答、检索增强生成、智能助手和知识服务等场景中得到广泛使用。随着模型参数规模和训练语料规模持续扩大，模型不仅学习到一般语言能力和事实知识，也可能记忆训练数据或后训练数据中的敏感个人信息。在攻击者构造直接询问、文本补全、角色扮演或上下文诱导提示时，模型可能输出手机号、邮箱、身份标识等敏感内容，带来隐私泄露风险。")
    add_para(doc, "现有防护手段通常包括提示词约束、安全对齐、输出检测过滤和拒答模板。这些方法能够在部分场景下降低泄露概率，但主要作用于输入输出层，并不一定改变模型内部对目标事实的记忆。若攻击提示绕过拒答模板或触发模型的补全倾向，敏感值仍可能被生成。因此，如何在不重新训练整个模型的情况下，直接修改模型对目标隐私知识的行为，是大语言模型安全与可信 AI 中具有实践价值的问题。")
    add_para(doc, "模型编辑（model editing）提供了一个可行切入点。ROME、MEMIT 等方法尝试直接定位并修改 Transformer 中存储事实关联的中间层参数，使模型对特定 subject-relation-object 事实产生新的输出。本文将该思想迁移到隐私知识清洗任务中：对于目标隐私事实，不再将其改写为另一个普通事实，而是将模型行为编辑为安全拒答；同时要求同一人物公开信息、其他人物公开信息和一般知识尽量保留。")
    heading(doc, "1.2  国内外研究现状", 2)
    add_para(doc, "知识编辑研究关注如何在不完全重训模型的情况下更新或删除局部知识。ROME 通过因果追踪定位与事实关联相关的中层 MLP 模块，并执行秩一权重更新；MEMIT 在 ROME 基础上扩展到多层分布式更新，支持批量修改多条知识；EasyEdit 则提供了统一模型编辑框架，便于在同一工程环境中运行多种编辑算法。")
    add_para(doc, "隐私保护方向的研究则包括训练数据提取攻击、机器遗忘、差分隐私训练、隐私神经元定位与后处理式清洗等路线。Carlini 等人的训练数据提取研究表明语言模型可能复现训练语料中的罕见序列；DEPN 等工作进一步尝试定位并编辑与隐私文本相关的神经元。与这些工作相比，本文不声称解决真实预训练 PII 删除问题，而是在可控 synthetic benchmark 上研究后训练注入隐私记忆的清洗流程、评价指标和副作用。")
    heading(doc, "1.3  本文主要工作与贡献", 2)
    add_numbered(doc, [
        "构建 private/public 对照的 synthetic privacy benchmark，使隐私值、公开事实和攻击提示均可控。",
        "实现 MLP-only LoRA 隐私记忆注入，构造能够稳定泄露目标隐私值的 merged leakage model。",
        "基于 EasyEdit 系统比较 ROME、MEMIT 与 PACE 在隐私压制、公开知识保持和过度拒答上的差异。",
        "完成指标定义审计，修正 Private Value Contains、PII-format Regex、Public Contains 等指标解释口径。",
        "提出 CAPE-v0 副作用感知请求选择方案，为后续降低 PACE 过度拒答提供实验设计基础。",
    ])
    heading(doc, "1.4  任务分工", 2)
    add_para(doc, "本项目为个人独立完成。主要工作包括数据构造、实验脚本开发、LoRA 注入训练、模型编辑实验、服务器运行管理、结果回收、指标审计、论文撰写与文档整理。AI 编程助手参与了部分代码生成、实验脚本整理、日志分析和文字润色，但研究设计、实验选择、结果解释和最终报告组织由本人完成并审核。")


def add_chapter_2(doc):
    heading(doc, "2  问题定义", 1)
    heading(doc, "2.1  相关背景知识", 2)
    add_para(doc, "自回归大语言模型可视为从输入序列到输出序列概率分布的函数。模型在生成时根据上下文预测下一个 token，事实知识和模式记忆可能分布在注意力层、MLP/FFN 层以及残差流中。既有研究表明，部分事实关联可通过中间层因果干预和权重更新进行定位与修改。")
    add_para(doc, "LoRA 是一种参数高效微调方法，通过在原始权重旁路中加入低秩矩阵实现任务适配。本文使用 MLP-only LoRA 将可控 private/public facts 注入 Qwen2.5-7B，使模型形成可评估的隐私记忆。模型编辑阶段不再重新训练 LoRA，而是使用 ROME/MEMIT 等编辑器将目标隐私事实的输出改为拒答。")
    heading(doc, "2.2  术语定义", 2)
    add_numbered(doc, [
        "隐私知识三元组：一条隐私知识记为 (e, a, v)，其中 e 表示人物实体，a 表示属性类型，v 表示隐私属性值。",
        "等价攻击提示集合：对同一隐私事实构造 direct、paraphrase、completion、roleplay、context 五类提示，用于检验编辑后的泛化效果。",
        "公开知识保持：同一人物公开事实、其他人物相同关系事实和一般知识问题在编辑后仍应尽量回答正确。",
        "过度拒答：模型在非隐私公开问题上也输出拒答文本，导致 public retain 下降。",
    ])
    heading(doc, "2.3  隐私知识清洗任务定义", 2)
    add_para(doc, "给定基础模型 M、待清洗隐私知识集合 K_F、需保留公开知识集合 K_R 和安全拒答集合 S，目标是通过模型编辑得到 M'，使 M' 在隐私攻击提示上尽量不输出目标隐私值或敏感格式，同时在公开事实查询上尽量保持原有回答能力。该任务与传统知识更新不同：目标输出不是另一个事实值，而是拒答行为；同时拒答本身也可能泛化到 public query，形成副作用。")
    heading(doc, "2.4  整体实验流程", 2)
    add_numbered(doc, [
        "构造 synthetic private/public facts 与攻击提示模板。",
        "使用 MLP-only LoRA 注入隐私记忆，得到 merged leakage model。",
        "使用 ROME/MEMIT 执行 direct-only 隐私拒答编辑。",
        "运行 full private eval 与 full public eval，审计 leakage、retain 与 refusal。",
        "使用 PACE 对 residual leakage 进行闭环再编辑，并进一步分析 over-refusal。",
        "基于副作用分析设计 CAPE-v0 请求选择策略。",
    ])
    heading(doc, "2.5  涉及的相关知识和技能", 2)
    add_para(doc, "项目涉及 Transformer 与 MLP 层机制、LoRA 参数高效微调、ROME/MEMIT 模型编辑、EasyEdit 工程框架、隐私泄露检测、正则表达式评估、批量生成评测、Linux/AutoDL 服务器实验管理、Git 协作与实验 artifact 管理等技术。")


def add_chapter_3(doc, data):
    heading(doc, "3  问题的主要解决方案", 1)
    heading(doc, "3.1  总体框架与设计动机", 2)
    add_para(doc, "本文方案不是单纯调用一个模型编辑算法，而是构建“隐私记忆注入—模型编辑清洗—多视角评估—失败样本再编辑—副作用分析”的完整流程。其核心动机在于：隐私场景中的攻击提示具有多样性，仅编辑 canonical QA prompt 很难保证对 completion、roleplay、context 等形式的泛化；同时拒答编辑容易误伤人物相关公开知识，因此必须同时评估 privacy suppression 与 public retention。")
    add_para(doc, "总体框架包括四个层次：第一层是 synthetic benchmark，用于产生可控隐私事实、公开事实和攻击提示；第二层是 LoRA 注入，用于构造具备明确泄露行为的待清洗模型；第三层是 ROME/MEMIT 基础编辑器；第四层是 PACE 与 CAPE 等围绕隐私任务设计的 wrapper/strategy。")
    add_para(doc, "图示预留：基于模型编辑的隐私知识清洗总体框架图可按“数据构造—隐私注入—模型编辑—评估审计—闭环改进”的流程绘制，图题建议为“图1 基于模型编辑的隐私知识清洗总体框架”。")
    heading(doc, "3.2  Synthetic Privacy Benchmark 构建", 2)
    add_para(doc, "Benchmark 以 100 个虚拟人物为主体，构造 private facts、same-subject public facts、same-relation other-subject public facts 和 general knowledge 问题。每条 private fact 配置五类攻击提示，每条 public fact 配置直接问答提示。该设计使隐私清洗不只考察目标值是否被压制，也考察公开事实和一般知识是否被保留。")
    add_table(doc, [
        ["虚拟人物", "100"],
        ["Private facts", "200"],
        ["Public facts", "840"],
        ["General knowledge", "40"],
        ["Private attack types", "direct / paraphrase / completion / roleplay / context"],
        ["Full private eval prompts", "3000"],
        ["Full public eval prompts", "2520"],
    ], ["数据项", "规模"])
    heading(doc, "3.3  MLP-only LoRA 隐私记忆注入", 2)
    add_para(doc, "为了避免直接使用真实个人信息，本文通过 LoRA 将 synthetic private/public facts 注入 Qwen2.5-7B。LoRA 作用范围限制在 MLP 模块（gate_proj、up_proj、down_proj），使模型在后训练阶段形成可控记忆。合并 adapter 后得到 merged leakage model，作为 ROME、MEMIT 与 PACE 的共同起点。")
    add_para(doc, "该步骤的作用是建立实验前提：如果模型本身没有稳定记住目标隐私值，则后续清洗实验无法判断编辑方法是否真正压制了已存在的隐私记忆。")
    heading(doc, "3.4  基于 ROME / MEMIT 的隐私拒答编辑", 2)
    add_para(doc, "ROME direct-only 以每个隐私事实的 canonical QA prompt 为编辑输入，将原始 ground truth 隐私值替换为统一拒答文本。MEMIT direct-only 使用相同数据、模型和请求，但采用 MEMIT 配置进行多层编辑。二者均不修改 EasyEdit 底层算法，仅通过不同编辑器比较隐私清洗效果。")
    add_para(doc, "需要强调的是，MEMIT 在当前实验中不是“更优方法”，而是提供了高 public retain、弱 privacy suppression 的对照端点。它帮助说明隐私压制和公开知识保持之间存在权衡。")
    heading(doc, "3.5  PACE 隐私感知闭环编辑策略", 2)
    add_para(doc, "PACE（Privacy-Aware Closed-loop Editing，隐私感知闭环编辑）不是新的底层模型编辑算法，而是基于 ROME/MEMIT 等编辑器之上的 residual leakage re-edit wrapper。其基本思想是在 direct-only 编辑后运行 full private eval，收集仍泄露或仍输出敏感格式的失败样本，再构造 round2 requests 追加编辑。")
    add_para(doc, "PACE 的价值在于检验 failure-driven re-edit 是否能扩大隐私压制覆盖面。实验结果表明，该策略确实能够显著降低泄露，但也容易造成 public query 上的高拒答率，因此需要进一步控制副作用。")
    heading(doc, "3.6  CAPE-v0 副作用感知请求选择策略", 2)
    cape = data["cape_report"]
    add_para(doc, "基于 PACE 的过度拒答现象，本文进一步设计 CAPE-v0（Collateral-Aware Privacy Editing，副作用感知隐私编辑）。CAPE-v0 同样是 request selection / wrapper，不改变 ROME/MEMIT 底层更新规则。其核心机制包括 residual leakage re-edit、public-anchor blocking 和 per-person edit budget。")
    add_para(doc, "CAPE-v0 的机制图可围绕 residual leakage candidates、public-anchor blocking、per-person budget 与 selected round2 requests 四个节点绘制，用于说明其如何在 naive PACE 与 ROME 之间探索更合理的 privacy-utility trade-off。")
    add_para(doc, f"当前已完成 CAPE-v0 的请求选择阶段。在 B=1、tau=0.5 设置下，脚本从 {cape['selection_summary']['num_candidates']} 个 residual leakage candidates 中选出 {cape['selection_summary']['num_selected']} 条 round2 re-edit requests；其中 public-anchor blocking 跳过 {cape['selection_summary']['skipped_public_anchor']} 条，per-person budget 跳过 {cape['selection_summary']['skipped_budget']} 条。该结果说明 CAPE-v0 能在进入模型编辑前显式控制潜在副作用，但其完整编辑效果仍需等待服务器 GPU 实验完成后补充。")
    add_table(doc, [
        ["Residual leakage candidates", cape["selection_summary"]["num_candidates"]],
        ["Public-anchor blocking skipped", cape["selection_summary"]["skipped_public_anchor"]],
        ["Per-person budget skipped", cape["selection_summary"]["skipped_budget"]],
        ["Selected round2 requests", cape["selection_summary"]["num_selected"]],
        ["Selected people", cape["selection_summary"]["num_selected_people"]],
        ["Selected attack type", json.dumps(cape["selection_summary"]["by_attack_type"], ensure_ascii=False)],
    ], ["CAPE-v0 选择项", "数值"])
    heading(doc, "3.7  方法复杂度与工程实现说明", 2)
    add_para(doc, "本文将高成本操作限制在 LoRA 注入、ROME/MEMIT 编辑和 full eval 阶段；数据构造、指标审计、请求选择和结果汇总均可在本地 CPU 环境完成。CAPE-v0 的请求选择阶段只读取 JSON/CSV artifacts，不加载大模型，因此适合在 Windows 本地生成小文件后通过 Git 同步到服务器运行。")


def add_chapter_4(doc, data):
    heading(doc, "4  实验验证结果和分析讨论", 1)
    mc = data["method_comparison"]
    over = data["over_refusal"]
    over_map = {r["method_key"]: r for r in over}
    heading(doc, "4.1  数据集和实验设置", 2)
    add_para(doc, "实验使用 Qwen2.5-7B 作为基础模型。首先通过 MLP-only LoRA 构造 merged leakage model，然后在同一模型起点上分别运行 ROME direct-only、MEMIT direct-only 和三组 PACE 变体。所有方法均使用同一 v2 synthetic dataset 和同一批 direct requests，保证模型、数据和评估口径可比。")
    heading(doc, "4.2  评价指标与审计说明", 2)
    add_para(doc, "指标审计发现，原脚本中的 private_exact 并不是严格 equality，而是 normalize(value) in normalize(output) 的目标值包含匹配。因此本文统一将其称为 Private Value Contains。PII-format Regex 只覆盖 phone/email 格式，不覆盖所有 private attributes；Sensitive Pattern 统计敏感格式输出，不要求与目标值一致；Public Contains 是宽松包含式公开知识保持指标，不等价于严格 factual accuracy；Private/Public Refusal 基于拒答模板匹配，需结合泄露指标共同解释。")
    heading(doc, "4.3  实验一：LoRA 隐私注入与泄露模型构建", 2)
    pre = next(r for r in mc if r["method_key"] == "pre_leakage_model")
    add_para(doc, f"合并 LoRA 后的 leakage model 在 full private eval 中表现出稳定泄露：Private Value Contains 为 {fmt(pre['private_exact'])}，PII-format Regex 为 {fmt(pre['private_regex'])}，Sensitive Pattern 为 {fmt(pre['private_sensitive'])}；同时 Public Contains 达到 {fmt(pre['public_contains'])}。这说明模型同时记住了 private facts 和 public facts，后续清洗实验具备成立前提。")
    heading(doc, "4.4  实验二：ROME direct-only 隐私编辑", 2)
    rome = next(r for r in mc if r["method_key"] == "rome_direct_only")
    add_para(doc, f"ROME direct-only 将 Private Value Contains 从 {fmt(pre['private_exact'])} 降至 {fmt(rome['private_exact'])}，Private Refusal 提升至 {fmt(rome['private_refusal'])}，表明 ROME 具备一定隐私压制能力。但 Public Contains 同时从 {fmt(pre['public_contains'])} 降至 {fmt(rome['public_contains'])}，Public Refusal 达到 {fmt(over_map['rome_direct_only']['public_refusal_rate'])}，说明其副作用不可忽视。")
    heading(doc, "4.5  实验三：MEMIT direct-only 隐私编辑", 2)
    memit = next(r for r in mc if r["method_key"] == "memit_direct_only")
    add_para(doc, f"MEMIT direct-only 的 Public Contains 为 {fmt(memit['public_contains'])}，高于 ROME 的 {fmt(rome['public_contains'])}；其 Public Refusal 为 {fmt(over_map['memit_direct_only']['public_refusal_rate'])}，低于 ROME。然而 MEMIT 的 Private Value Contains 仍为 {fmt(memit['private_exact'])}，Private Refusal 仅为 {fmt(memit['private_refusal'])}。因此，MEMIT 在当前拒答编辑任务中主要体现为较高公开知识保持但隐私压制不足，不能解释为整体优于 ROME。")
    heading(doc, "4.6  实验四：PACE 闭环再编辑", 2)
    pace2 = next(r for r in mc if r["method_key"] == "pace_max2_per_person")
    add_para(doc, f"PACE 通过对 residual leakage 进行再编辑，能够显著降低隐私泄露。以 max2_per_person 为例，Private Value Contains 降至 {fmt(pace2['private_exact'])}，PII-format Regex 降至 {fmt(pace2['private_regex'])}，Sensitive Pattern 降至 {fmt(pace2['private_sensitive'])}。但其 Public Contains 仅为 {fmt(pace2['public_contains'])}，Public Refusal 高达 {fmt(over_map['pace_max2_per_person']['public_refusal_rate'])}。这表明 naive residual re-edit 在提高隐私压制的同时，将大量人物相关公开问题推向拒答。")
    heading(doc, "4.7  主结果横向对比", 2)
    rows = []
    for r in mc:
        o = over_map[r["method_key"]]
        rows.append([r["method"], fmt(r["private_exact"]), fmt(r["private_regex"]), fmt(r["private_sensitive"]), fmt(r["private_refusal"]), fmt(r["public_contains"]), fmt(o["public_refusal_rate"])])
    add_table(doc, rows, ["方法", "Private Value Contains↓", "PII-format Regex↓", "Sensitive Pattern↓", "Private Refusal↑", "Public Contains↑", "Public Refusal↓"])
    add_para(doc, "主结果表明，不同策略形成了清晰的权衡结构：ROME 位于中等隐私压制与中等公开知识损伤之间；MEMIT 位于较高公开知识保持但弱隐私压制一侧；PACE 位于强隐私压制但强副作用一侧。")
    heading(doc, "4.8  Over-refusal 分析", 2)
    add_table(doc, [[r["method"], r["public_refusal_count"], r["public_refusal_denominator"], fmt(r["public_refusal_rate"]), fmt(r["public_contains"])] for r in over], ["方法", "Public Refusal Count", "Denominator", "Public Refusal Rate↓", "Public Contains↑"])
    add_para(doc, "过度拒答分析说明，public retain 下降不只是回答错误，还包含明显的公开问题拒答。尤其是 PACE target_only 和 max1_per_case 的 Public Refusal 均超过 0.96，max2_per_person 也达到 0.8115。这一现象提示：高拒答率不能直接等价为隐私清洗成功，必须同时报告 public refusal。")
    heading(doc, "4.9  Attack-type 分析", 2)
    atk = data["attack_type"]
    selected_methods = ["rome_direct_only", "memit_direct_only", "pace_max2_per_person"]
    rows = []
    for r in atk:
        if r["method_key"] in selected_methods:
            rows.append([r["method"], r["attack_type"], fmt(r["private_exact"]), fmt(r["private_sensitive"]), fmt(r["private_refusal"])])
    add_table(doc, rows[:20], ["方法", "Attack Type", "Private Value Contains↓", "Sensitive Pattern↓", "Private Refusal↑"])
    add_para(doc, "attack-type 拆分显示，completion/context 等非直接问答形式对 direct-only editing 构成更强挑战。该结果支持本文对闭环再编辑和副作用感知请求选择的设计动机。")
    heading(doc, "4.10  Privacy-Utility Trade-off 分析", 2)
    trade = data["tradeoff"]
    add_table(doc, [[r["method"], fmt(r["privacy_score"]), fmt(r["utility_score"]), fmt(r["pus"])] for r in trade], ["方法", "PrivacyScore=1-PVC", "UtilityScore=Public Contains", "PUS"])
    if TRADEOFF_PNG.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(TRADEOFF_PNG), width=Inches(5.5))
        cap = doc.add_paragraph("图1 不同编辑策略下的隐私压制—公开知识保持权衡")
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_para(doc, "PrivacyScore 定义为 1 - Private Value Contains，UtilityScore 定义为 Public Contains。PUS 仅作为简化综合指标，用于展示 trade-off 趋势。结果显示，PACE 的 PrivacyScore 最高但 UtilityScore 极低；MEMIT 的 UtilityScore 较高但 PrivacyScore 较低；ROME 处于二者之间。")
    add_para(doc, "上述结果共同表明，当前任务的核心难点不是单独降低目标隐私输出，而是在 privacy-utility trade-off 中同时控制隐私泄露、公开知识保持与过度拒答。")
    heading(doc, "4.11  典型案例分析", 2)
    add_para(doc, "定性样例显示，部分编辑后输出会同时包含拒答模板和敏感格式，说明 safe_refusal 不能单独作为成功指标；另一些 public query 中模型直接拒答，说明人物相关问题会受到隐私编辑的泛化影响。该现象与 over-refusal 统计相互印证。")
    heading(doc, "4.12  实验小结", 2)
    add_para(doc, "本章实验说明，模型编辑可用于降低 synthetic privacy setting 下的目标隐私输出，但当前策略尚未同时兼顾隐私压制与公开知识保持。该结论不是对方法价值的否定，而是为副作用感知编辑策略提供了明确实验依据。")


def add_chapter_5(doc):
    heading(doc, "5  代码/软件使用说明", 1)
    heading(doc, "5.1  项目目录结构", 2)
    add_para(doc, "项目基于 EasyEdit 仓库扩展，主要目录包括 scripts/、docs/、artifacts/、hparams/、easyeditor/ 等。其中 scripts/ 存放数据构造、LoRA 训练、模型编辑、评估、汇总和 CAPE 请求选择脚本；artifacts/ 存放每轮实验的小型结果文件和审计报告；hparams/ 存放 ROME/MEMIT 等方法配置；easyeditor/ 为 EasyEdit 原始框架主体。")
    heading(doc, "5.2  数据构造模块", 2)
    add_para(doc, "数据构造脚本负责生成 synthetic_privacy_dataset.json、LoRA 训练数据和 ROME/MEMIT 请求文件。每条样本保留 case_id、person_id、attribute、attack_type、attack_template_id 等字段，保证后续评估能够按攻击类型、人物和属性进行追踪。")
    heading(doc, "5.3  LoRA 隐私注入模块", 2)
    add_para(doc, "LoRA 注入模块使用 EasyEdit/Transformers/PEFT 相关能力，对 Qwen2.5-7B 进行 MLP-only 训练，并将 adapter merge 为独立模型目录。训练结束后通过 private/public eval 验证模型是否形成可控泄露。")
    heading(doc, "5.4  ROME / MEMIT 编辑模块", 2)
    add_para(doc, "隐私拒答编辑统一使用 scripts/run_privacy_refusal_edit.py。该脚本支持 ROME 和 MEMIT，能够加载外部 requests_path，并在编辑完成后自动生成 subset/full private eval 和 public retain eval。MEMIT 所需 MOM2 stats 通过独立准备脚本生成，不关闭 mom2_adjustment，也不修改底层 MEMIT 算法。")
    heading(doc, "5.5  PACE / CAPE 请求选择模块", 2)
    add_para(doc, "PACE 请求构造基于 ROME 后 residual leakage cases，支持 target_only、max1_per_case、max2_per_person 等不同选择约束。CAPE-v0 则进一步读取 public retain eval，加入 public-anchor blocking 和 per-person budget，使请求选择在进入 GPU 编辑前就体现副作用控制。")
    heading(doc, "5.6  评估与审计模块", 2)
    add_para(doc, "评估模块包括 evaluate_privacy_leakage.py、evaluate_public_retain.py、compare_v2_editing_methods.py 和 analyze_v2_audit.py。审计脚本会输出 metric definitions、method provenance、over-refusal、attack-type breakdown、trade-off points 和 qualitative cases，避免只凭单一泄露率解释结果。")
    heading(doc, "5.7  实验复现流程", 2)
    add_numbered(doc, [
        "在本地修改脚本并通过 git push 同步到 GitHub。",
        "在 AutoDL 服务器进入 /root/autodl-tmp/projects/EasyEdit 后 git pull。",
        "激活 conda easyedit 环境，设置 HF cache、PYTHONPATH、代理与 NLTK_DATA。",
        "运行对应 pipeline 脚本完成编辑、评估、汇总和 artifact 打包。",
        "仅将 summary/eval/csv/md/png 等小文件拉回本地分析，不传输模型权重、MOM2 stats 或 Wikipedia cache。",
    ])
    heading(doc, "5.8  运行环境与依赖", 2)
    add_para(doc, "主要运行环境为 Linux/AutoDL GPU 服务器，模型与缓存位于 /root/autodl-tmp。本地 Windows 负责代码修改、结果分析和论文整理。关键依赖包括 PyTorch、Transformers、Datasets、EasyEdit、PEFT、python-docx 等。")
    heading(doc, "5.9  AI 辅助开发记录", 2)
    add_para(doc, "项目开发过程中使用 AI 编程助手辅助生成脚本、整理运行命令、分析日志和润色论文文本。所有实验命令、结果口径和最终结论均经过人工确认，AI 作为辅助工具而非独立研究主体。")


def add_chapter_6_and_refs(doc):
    heading(doc, "6  结论", 1)
    add_para(doc, "本文围绕大语言模型隐私知识清洗任务，构建了从 synthetic privacy benchmark、LoRA 隐私记忆注入、ROME/MEMIT 模型编辑、PACE 闭环再编辑到指标审计和副作用分析的完整研究流程。实验验证了模型编辑压制可控隐私泄露的可行性，也揭示了当前策略在隐私压制和公开知识保持之间存在明显权衡。")
    add_para(doc, "具体而言，ROME 能降低部分目标隐私值输出，但会带来明显 public damage；MEMIT 保留公开知识较好，但隐私压制能力不足；PACE 能显著压低泄露，却容易导致人物相关公开问题上的过度拒答。这些结果说明，隐私知识清洗不应只追求更高拒答率或更低目标值包含率，还必须同时控制 non-target knowledge 的副作用。")
    add_para(doc, "基于上述发现，本文提出 CAPE-v0 作为后续改进方向，通过 public-anchor blocking 与 per-person edit budget 在请求选择阶段控制副作用。当前 CAPE-v0 已完成请求选择和报告生成，完整模型编辑评估仍需等待服务器实验结果回收后补充。本文的局限包括：实验基于 synthetic privacy setting，不代表真实预训练 PII 删除；PACE/CAPE 的 residual mining 与最终评估仍存在 held-out separation 不充分的问题；Public Contains 属于宽松保持指标，后续仍需更严格的人工或语义评估。")
    heading(doc, "参考文献", 1)
    refs = [
        "Carlini N., Tramer F., Wallace E., et al. Extracting Training Data from Large Language Models. USENIX Security Symposium, 2021.",
        "Meng K., Bau D., Andonian A., Belinkov Y. Locating and Editing Factual Associations in GPT. Advances in Neural Information Processing Systems, 2022.",
        "Meng K., Sharma A. S., Andonian A., Belinkov Y., Bau D. Mass-Editing Memory in a Transformer. arXiv:2210.07229, 2022.",
        "Wang P., Zhang N., Tian B., et al. EasyEdit: An Easy-to-use Knowledge Editing Framework for Large Language Models. ACL Demo, 2024.",
        "Hu E. J., Shen Y., Wallis P., et al. LoRA: Low-Rank Adaptation of Large Language Models. ICLR, 2022.",
        "Bourtoule L., Chandrasekaran V., Choquette-Choo C. A., et al. Machine Unlearning. IEEE Symposium on Security and Privacy, 2021.",
        "Wu X., Li J., Xu M., et al. DEPN: Detecting and Editing Privacy Neurons in Pretrained Language Models. EMNLP, 2023.",
        "Dai D., Dong L., Hao Y., Sui Z., Chang B., Wei F. Knowledge Neurons in Pretrained Transformers. ACL, 2022.",
        "Petroni F., Rocktaschel T., Riedel S., et al. Language Models as Knowledge Bases? EMNLP-IJCNLP, 2019.",
        "Brown H., Lee K., Mireshghallah F., Tramer F., Shokri R. What Does It Mean for a Language Model to Preserve Privacy? ACM FAccT, 2022.",
        "Lukas N., Salem A., Sim R., Tople S., Wutschitz L., Zanella-Beguelin S. Analyzing Leakage of Personally Identifiable Information in Language Models. IEEE Symposium on Security and Privacy, 2023.",
        "Dwork C., Roth A. The Algorithmic Foundations of Differential Privacy. Foundations and Trends in Theoretical Computer Science, 2014.",
        "Jang J., Yoon D., Yang S., et al. Knowledge Unlearning for Mitigating Privacy Risks in Language Models. arXiv preprint, 2023.",
        "Eldan R., Russinovich M. Who's Harry Potter? Approximate Unlearning in LLMs. arXiv:2310.02238, 2023.",
    ]
    for i, ref in enumerate(refs, 1):
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Pt(0)
        p.add_run(f"[{i}] {ref}")


def build_docx(data):
    doc = Document()
    set_doc_styles(doc)
    add_cover(doc)
    add_abstract(doc)
    add_chapter_1(doc)
    add_chapter_2(doc)
    add_chapter_3(doc, data)
    add_chapter_4(doc, data)
    add_chapter_5(doc)
    add_chapter_6_and_refs(doc)
    doc.save(OUT_DOCX)


def write_log(data):
    cape = data["cape_report"]["selection_summary"]
    lines = [
        "# COURSE PAPER REWRITE LOG",
        "",
        "## 1. Read Files",
        "",
        f"- `{SOURCE_DOCX.name}`: original draft, backed up before rewrite.",
        f"- `{TEMPLATE_DOCX.name}`: read-only course paper template for structure reference.",
        "- `artifacts/analysis_v2_audit_20260622/`: metric audit, method comparison, over-refusal, attack-type and trade-off outputs.",
        "- `artifacts/run_20260622_v2_memit_direct/`: MEMIT direct-only result artifacts.",
        "- `artifacts/run_20260622_v2_cape_b1_tau05/`: CAPE-v0 request-selection report.",
        "",
        "## 2. Updated Sections",
        "",
        "- Rewritten abstract, keywords, introduction, problem definition, method, experiments, software usage and conclusion.",
        "- Updated stale statements about missing MEMIT; MEMIT direct-only is now included as completed baseline.",
        "- Unified metric names: `Private Value Contains`, `PII-format Regex`, `Sensitive Pattern`, `Public Contains`, and `Public Refusal`.",
        "- Unified PACE naming as `Privacy-Aware Closed-loop Editing`.",
        "- Added CAPE-v0 as follow-up side-effect-aware request selection strategy.",
        "",
        "## 3. Added Experimental Results",
        "",
        "- Added main ROME / MEMIT / PACE comparison table.",
        "- Added public over-refusal analysis.",
        "- Added privacy-utility trade-off section and inserted `privacy_utility_tradeoff.png`.",
        "- Added CAPE-v0 request-selection status.",
        f"- CAPE-v0 current selection: candidates={cape['num_candidates']}, selected={cape['num_selected']}, public-anchor skipped={cape['skipped_public_anchor']}, budget skipped={cape['skipped_budget']}.",
        "",
        "## 4. Pending Server Results",
        "",
        "- CAPE-v0 model editing and full private/public eval are not yet written as completed results.",
        "- After server returns CAPE eval JSON, update Chapter 4 tables, trade-off plot and conclusion.",
        "",
        "## 5. Figures Requiring Manual Check",
        "",
        "- Existing `privacy_utility_tradeoff.png` has been inserted into the Word document.",
        "- Overall framework figure and CAPE-v0 mechanism figure are described in text but not yet drawn as independent figures.",
        "- LibreOffice/soffice was not found in the current Windows environment, so automated render-to-PNG visual QA could not be completed here. Please open the final DOCX in Word/WPS for final visual layout inspection.",
        "",
        "## 6. References to Verify",
        "",
        "- Reference list uses known model editing, privacy leakage, LoRA, EasyEdit, DEPN, machine unlearning and differential privacy literature.",
        "- DOI fields are intentionally omitted to avoid fabricating identifiers; final submission can add DOI after manual verification.",
        "",
        "## 7. Self Check",
        "",
        "- Template file was not modified.",
        "- No CAPE model-editing result was fabricated.",
        "- `Private Value Contains` is not described as strict equality.",
        "- `Public Contains` is not described as strict factual accuracy.",
        "- MEMIT direct-only is included.",
        "- PACE is described as a wrapper/strategy, not a bottom-layer editing algorithm.",
    ]
    LOG_MD.write_text("\n".join(lines) + "\n", encoding="utf-8")


def main():
    if SOURCE_DOCX.exists() and not BACKUP_DOCX.exists():
        shutil.copy2(SOURCE_DOCX, BACKUP_DOCX)

    data = {
        "method_comparison": read_csv(ANALYSIS / "method_comparison.csv"),
        "over_refusal": read_csv(ANALYSIS / "over_refusal_stats.csv"),
        "attack_type": read_csv(ANALYSIS / "attack_type_breakdown.csv"),
        "tradeoff": read_csv(ANALYSIS / "tradeoff_points.csv"),
        "cape_report": read_json(CAPE_DIR / "cape_selection_report.json"),
    }
    DRAFT_MD.write_text(build_markdown(data), encoding="utf-8")
    build_docx(data)
    write_log(data)
    print(OUT_DOCX)
    print(DRAFT_MD)
    print(LOG_MD)


if __name__ == "__main__":
    main()
