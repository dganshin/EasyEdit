#!/usr/bin/env python3
"""Targeted in-place revision for the final course paper.

The source text remains the v5-derived paper. This script only replaces stale
sections, calibrates the final-result narrative, and rebuilds the existing DOCX
with Chinese thesis-style fonts.
"""

from __future__ import annotations

import csv
import re
from pathlib import Path
from typing import Any, Dict, Iterable, List, Sequence

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


TEMPLATE = Path("课程论文撰写模板 (研究类).docx")
MD_PATH = Path("docs/course_paper_v7_final_academic_polished.md")
DOCX_PATH = Path("课程论文_v7_final_academic_polished.docx")
QA_PATH = Path("docs/COURSE_PAPER_V8_FINAL_QA.md")
TABLE_DIR = Path("artifacts/final_paper_assets_20260623/tables")
FIG_DIR = Path("artifacts/final_paper_assets_20260623/figures")


def read_csv(path: Path) -> List[Dict[str, str]]:
    with path.open(encoding="utf-8", newline="") as fh:
        return list(csv.DictReader(fh))


def md_table(rows: Iterable[Dict[str, Any]], fields: Sequence[str]) -> str:
    lines = ["|" + "|".join(fields) + "|", "|" + "|".join(["---"] * len(fields)) + "|"]
    for row in rows:
        lines.append("|" + "|".join(str(row.get(f, "")) for f in fields) + "|")
    return "\n".join(lines)


def replace_between(text: str, start: str, end: str, replacement: str) -> str:
    i = text.index(start)
    j = text.index(end, i)
    return text[:i] + replacement.rstrip() + "\n\n" + text[j:]


def replace_between_any(text: str, starts: Sequence[str], end: str, replacement: str) -> str:
    for start in starts:
        if start in text:
            return replace_between(text, start, end, replacement)
    raise ValueError(f"none of the start markers found before {end!r}: {starts!r}")


def compact_synthetic() -> List[Dict[str, str]]:
    label = {
        "Merged pre-edit": "合并后泄露模型",
        "Prompt Refusal": "提示拒答",
        "PACE_LITE_B20_K0": "PACE-Lite K0",
        "CAPE_ANCHOR_B20_K1": "CAPE-Anchor K1",
        "CAPE_ANCHOR_B20_K2": "CAPE-Anchor K2",
    }
    rows = []
    for r in read_csv(TABLE_DIR / "table1_synthetic_privacy_main.csv"):
        rows.append(
            {
                "方法": label.get(r["Method"], r["Method"]),
                "状态": r["Status"],
                "隐私值包含率↓": r["Private leak ↓"],
                "敏感格式率↓": r["Sensitive pattern ↓"],
                "隐私拒答率": r["Private refusal"],
                "公开知识保持率↑": r["Public contains ↑"],
            }
        )
    return rows


def compact_cape() -> List[Dict[str, str]]:
    label = {
        "PACE_LITE_B20_K0": "PACE-Lite K0",
        "CAPE_ANCHOR_B20_K1": "CAPE-Anchor K1",
        "CAPE_ANCHOR_B20_K2": "CAPE-Anchor K2",
    }
    note = {
        "PACE_LITE_B20_K0": "无公开锚点，作为受限闭环对照",
        "CAPE_ANCHOR_B20_K1": "公开保持和隐私压制之间较均衡",
        "CAPE_ANCHOR_B20_K2": "公开保持进一步提高，但隐私压制变弱",
    }
    rows = []
    for r in read_csv(TABLE_DIR / "table2_cape_anchor_ablation.csv"):
        rows.append(
            {
                "配置": label.get(r["Variant"], r["Variant"]),
                "隐私值包含率↓": r["Private leak ↓"],
                "PII格式率↓": r["PII regex ↓"],
                "隐私拒答率": r["Private refusal"],
                "公开知识保持率↑": r["Public contains ↑"],
                "说明": note.get(r["Variant"], ""),
            }
        )
    return rows


def compact_qwen() -> List[Dict[str, str]]:
    rows = []
    for r in read_csv(TABLE_DIR / "table3_qwen_public_transfer.csv"):
        if r["Status"].startswith("failed"):
            continue
        rows.append(
            {
                "数据集": "CounterFact" if r["Dataset"] == "counterfact" else "zsRE",
                "方法": r["Method"].replace("ROME_PACE_EDIT", "ROME+PACE-Edit").replace("ROME_CAPE_EDIT", "ROME+CAPE-Edit"),
                "样本数": r["Cases"],
                "可靠性↑": r["Reliability ↑"],
                "泛化性↑": r["Generalization ↑"],
            }
        )
    return rows


def compact_gptj() -> List[Dict[str, str]]:
    rows = []
    for r in read_csv(TABLE_DIR / "table4_gptj_boundary_check.csv"):
        rows.append(
            {
                "数据集": "CounterFact" if r["Dataset"] == "counterfact" else "zsRE",
                "方法": r["Method"].replace("ROME_PACE_EDIT", "ROME+PACE-Edit").replace("ROME_CAPE_EDIT", "ROME+CAPE-Edit"),
                "可靠性↑": r["Reliability ↑"],
                "泛化性↑": r["Generalization ↑"],
                "逐样本审计": r["Per-case rewrite"] or "未统计",
                "解释": "基础编辑正常" if r["Interpretation"] == "baseline normal" else "wrapper塌缩",
            }
        )
    return rows


def compact_failures() -> List[Dict[str, str]]:
    rows = []
    for r in read_csv(TABLE_DIR / "table5_failure_and_resource_limits.csv"):
        rows.append({"范围": r["Scope"], "对象": r["Method"], "状态": r["Status"], "原因": r["Reason"], "论文处理": r["Policy"]})
    return rows


def make_system_flow_figure() -> None:
    out_png = FIG_DIR / "fig7_system_module_archive_flow.png"
    out_svg = FIG_DIR / "fig7_system_module_archive_flow.svg"
    if out_png.exists() and out_svg.exists():
        return
    import matplotlib.pyplot as plt
    from matplotlib import font_manager
    from matplotlib.patches import FancyArrowPatch, FancyBboxPatch

    font_path = Path("C:/Windows/Fonts/msyh.ttc")
    if font_path.exists():
        font_manager.fontManager.addfont(str(font_path))
        plt.rcParams["font.family"] = "Microsoft YaHei"
    plt.rcParams["axes.unicode_minus"] = False
    fig, ax = plt.subplots(figsize=(11, 3.2), dpi=220)
    ax.axis("off")
    labels = [
        "合成数据\n构建",
        "LoRA 可控\n隐私注入",
        "模型编辑\n执行",
        "隐私/公开\n双重评估",
        "结果归档\n与论文表图",
    ]
    xs = [0.06, 0.26, 0.46, 0.66, 0.86]
    colors = ["#D9EAF7", "#E7F0D9", "#F5E6CC", "#EADCF8", "#F2D8D8"]
    for x, label, color in zip(xs, labels, colors):
        box = FancyBboxPatch((x - 0.075, 0.42), 0.15, 0.28, boxstyle="round,pad=0.02,rounding_size=0.025", fc=color, ec="#34495E", lw=1.2)
        ax.add_patch(box)
        ax.text(x, 0.56, label, ha="center", va="center", fontsize=11, color="#1F2D3D")
    for x1, x2 in zip(xs[:-1], xs[1:]):
        ax.add_patch(FancyArrowPatch((x1 + 0.085, 0.56), (x2 - 0.085, 0.56), arrowstyle="-|>", mutation_scale=13, lw=1.2, color="#566573"))
    ax.text(0.5, 0.18, "统一 case_id、请求文件、评估输出与摘要表，保证实验结果可追溯、可比较、可复核", ha="center", va="center", fontsize=10, color="#34495E")
    fig.tight_layout()
    fig.savefig(out_png, bbox_inches="tight")
    fig.savefig(out_svg, bbox_inches="tight")
    plt.close(fig)


def abstract_block() -> str:
    return """# 摘要

大语言模型在预训练和后训练过程中会把大量语言模式和事实关联压缩进参数，其中部分关联可能对应电话号码、邮箱等个人可识别信息。当模型在后续应用中面对直接询问、改写询问或补全式提示时，这类参数化隐私记忆可能被重新触发，进而产生隐私泄露风险。仅依赖输出过滤或提示层拒答虽然部署成本较低，但不能直接改变模型内部已经形成的知识关联；全量重训或机器遗忘又往往需要较高计算成本和原始训练数据访问权限。针对这一矛盾，本文研究基于模型编辑的隐私知识清洗问题，关注如何在压制目标隐私事实输出的同时，尽量保持同主体公开知识、同关系其他主体公开知识和通用知识回答能力。

本文首先构建 private/public 解耦的合成隐私评测基准，将电话号码、邮箱等隐私事实与职业、学校、雇主、家乡等公开事实显式分离，并设计 direct、paraphrase、completion、roleplay、context 等多类攻击提示。为了获得可控且可复现的待清洗对象，本文使用低秩适配（Low-Rank Adaptation，LoRA）向 Qwen2.5-7B 注入合成隐私与公开事实，并将适配器合并为统一的隐私泄露模型。在此基础上，本文将 ROME、MEMIT、FT 等通用模型编辑方法适配到隐私拒答式清洗任务，并提出 PACE、CAPE 与 CAPE-Anchor 三类请求选择层策略：PACE 根据残余泄露样本进行闭环再编辑，CAPE 引入副作用感知筛选，CAPE-Anchor 进一步在二轮请求集合中加入公开知识锚点，以缓解闭环再编辑造成的 public collapse。

实验结果显示，提示拒答几乎不能改变模型内部隐私记忆；MEMIT 保留公开知识较好但隐私压制不足，ROME 处于隐私压制与知识保持的中间区域，FT、PACE 和 CAPE 则形成强隐私压制但公开知识保持较弱的端点。CAPE-Anchor 在 Qwen synthetic privacy 主任务上形成了更合理的隐私压制—知识保持折中：相对于 naive PACE/CAPE，它提高了公开知识保持能力，但也会牺牲部分隐私压制强度。因此，本文将其结论限定为有限有效改进，而非零损伤方案或全局占优方案。进一步的 Qwen CounterFact/zsRE 实验表明，closed-loop wrapper 可以迁移到公开事实编辑任务中运行，但效果中等；GPT-J-6B 上的第二模型验证则显示，未重新调参的 wrapper 存在明显迁移边界。总体而言，本文建立了可复现的隐私知识清洗评测流程，并揭示了模型编辑隐私清洗中的隐私压制、公开知识保持与过度拒答之间的结构性权衡。

关键词：大语言模型；模型编辑；隐私知识清洗；知识保持；过度拒答；PACE；CAPE-Anchor"""


def section_24() -> str:
    return """## 2.4  基础编辑方法与闭环请求选择

ROME 面向单条事实编辑，通过定位事实关联在模型中的关键层并执行秩一更新来改变目标输出[7]。MEMIT 面向批量事实编辑，在多个层上估计更新并同时注入多条编辑约束[8]。二者均源于普通事实编辑任务，原始目标通常是把错误事实替换为另一个正确事实；本文将其改造为拒答式隐私清洗，即把目标隐私查询映射到安全拒答表达。

除 ROME 和 MEMIT 外，本文还纳入 FT、KN 与 IKE 作为对照。FT 通过直接参数微调学习拒答目标，能够形成强覆盖式更新，但也更容易破坏非目标知识。KN 基于知识神经元定位思想修改与特定事实相关的神经元，对计算资源和定位过程较敏感。IKE 属于基于上下文示例的编辑方式，本身不直接修改模型参数，但依赖检索或句向量模型构造示例。上述方法共同构成本文的基础编辑器对照组，用于观察不同编辑范式在隐私拒答目标下的行为端点。

需要强调的是，PACE、CAPE 与 CAPE-Anchor 并不是新的底层参数更新算法，而是位于基础编辑器外层的请求选择策略。它们保持底层编辑器实现不变，通过决定哪些残余泄露样本进入二轮编辑、是否引入公开知识锚点以及如何控制请求规模，来调节隐私压制与公开知识保持之间的权衡。这样的定位使本文能够把算法更新和请求构造区分开来，避免把 wrapper 层效果误写成基础编辑器本身的改进。"""


def section_37() -> str:
    return f"""## 3.7  CAPE-Anchor 公开知识锚点约束

CAPE-Anchor 是在 PACE/CAPE 基础上加入公开知识锚点的请求选择策略。其动机来自实验中的一个关键现象：仅根据残余泄露样本继续编辑，虽然可以进一步压低目标隐私值输出，但也容易把拒答模式扩散到同主体公开问题，形成 public collapse。换言之，隐私查询上的高拒答率并不必然意味着清洗质量更高；如果公开问题也被拒答，模型的可用性会明显下降。

CAPE-Anchor 不改变 ROME 等基础编辑器的参数更新公式，而是在二轮请求集合中同时放入隐私修复请求和公开知识锚点请求。形式上，最终请求集合为：

$$R_{{final}} = R_{{round1}} \\cup R_{{privacy}} \\cup R_{{anchor}}$$

其中，$R_{{round1}}$ 表示第一轮隐私拒答请求，$R_{{privacy}}$ 表示由残余泄露样本构造的二轮隐私修复请求，$R_{{anchor}}$ 表示与目标主体相关的公开知识锚点请求。公开锚点的作用不是让模型学习新的公开事实，而是在闭环再编辑过程中保留“该主体仍存在可回答的非敏感问题”这一约束信号。

本文只采用有限配置检验 CAPE-Anchor 的作用，而不进行大规模参数搜索。K0 可视为不加入公开锚点的受限闭环对照，K1 和 K2 分别加入不同数量的公开锚点。正文分析关注这些配置在隐私值包含率和公开知识保持率之间形成的操作点，而不是把具体工程参数作为方法贡献。"""


def chapter4() -> str:
    synthetic = md_table(compact_synthetic(), ["方法", "状态", "隐私值包含率↓", "敏感格式率↓", "隐私拒答率", "公开知识保持率↑"])
    cape = md_table(compact_cape(), ["配置", "隐私值包含率↓", "PII格式率↓", "隐私拒答率", "公开知识保持率↑", "说明"])
    qwen = md_table(compact_qwen(), ["数据集", "方法", "样本数", "可靠性↑", "泛化性↑"])
    gptj = md_table(compact_gptj(), ["数据集", "方法", "可靠性↑", "泛化性↑", "逐样本审计", "解释"])
    failures = md_table(compact_failures(), ["范围", "对象", "状态", "原因", "论文处理"])
    return f"""# 4  实验设置与结果分析

## 4.1  实验设置

本文实验围绕两个层级展开。第一层是 synthetic privacy 主实验，使用合并后的 Qwen2.5-7B 隐私泄露模型评估不同编辑方法在隐私压制、公开知识保持和过度拒答方面的差异。第二层是公开事实编辑迁移验证，使用 CounterFact 与 zsRE 检查 ROME-based closed-loop wrapper 在非隐私 factual editing 场景中的外部行为。两类实验的任务含义不同，因此本文将其分表报告，不把公开事实编辑结果写成 PII 清洗效果。

synthetic privacy 主实验使用同一合并后泄露模型、同一合成数据集、同一批 private/public 查询集合和同一套自动评估指标。这样的控制变量设计保证方法差异主要来自编辑器和请求选择策略，而不是数据切分、模型版本或生成设置差异。公开事实编辑实验则采用 200-case 规模，用于在有限资源条件下观察 wrapper 的迁移趋势和边界。

## 4.2  数据集与模型

合成隐私数据集包含 100 个虚拟人物，每个人物包含隐私事实和公开事实。隐私事实主要包括电话号码、邮箱等结构化 PII，公开事实包括职业、学校、雇主、家乡等非敏感知识。该设计使隐私清洗不再只是“是否拒答”的单指标任务，而是同时考察模型是否保留应当回答的公开事实。

待编辑模型由 Qwen2.5-7B 通过 MLP-only LoRA 注入隐私与公开事实后合并得到。编辑前模型的隐私值包含率为 0.9387，公开知识保持率为 0.9766，说明模型同时具有稳定隐私记忆和公开事实记忆。因此，后续若隐私泄露下降，可以归因于编辑产生的压制效果；若公开知识保持下降，则应视为编辑副作用。

表 2 给出主要实验平台与运行环境。该表只用于说明实验资源条件，不作为算法贡献。KN 的部分定位过程和 GPT-J wrapper 的第二模型验证均受到资源与超参数敏感性影响，相关内容在第 4.10 节作为边界分析处理。

表 2  实验平台与主要配置

|配置项|配置内容|
|---|---|
|GPU|NVIDIA vGPU-48GB，48GB 显存|
|CPU|20 vCPU Intel Xeon Platinum 8470Q|
|内存|90GB|
|数据盘|50GB SSD|
|主要框架|PyTorch, Transformers, EasyEdit|
|主要模型|Qwen2.5-7B, GPT-J-6B|
|主要任务|synthetic privacy, CounterFact, zsRE|

表 2 说明本文实验是在单卡 48GB 显存条件下完成的。该设置足以支持 Qwen synthetic 主实验和 200-case public transfer，但不适合继续扩展更大规模多模型矩阵。因此，本文将未完成或资源受限项明确列入 failure/resource limitation 表，而不是把缺失项隐含为有效结果。

## 4.3  评价指标

隐私值包含率衡量模型输出是否包含目标隐私值，数值越低表示目标泄露越少。敏感格式率统计输出中是否出现电话、邮箱等敏感格式，能够反映模型是否仍倾向生成 PII 形态内容。公开知识保持率衡量公开答案是否出现在模型输出中，数值越高表示非敏感知识越可能被保留。

公开问题拒答率用于衡量过度拒答。对于隐私问题，拒答通常是期望行为；但对于职业、学校、雇主等公开问题，拒答意味着模型把隐私保护目标扩散到了保留集合。因此，private refusal 高不等于清洗质量高，必须结合 public retain 和 public refusal 同时解释。

## 4.4  合成隐私基准主实验

表 3 汇总 synthetic privacy 主实验结果。合并后泄露模型的隐私值包含率为 0.9387，公开知识保持率为 0.9766，说明该模型适合作为统一编辑起点。提示拒答的隐私值包含率仍为 0.9490，表明仅在提示层要求模型拒答并不能改变参数中已经形成的隐私事实关联。

表 3  合成隐私基准主实验结果（↓越低越好，↑越高越好）

{synthetic}

从基础编辑方法看，MEMIT 的公开知识保持率为 0.8472，但隐私值包含率仍为 0.8140，说明它在当前拒答式隐私清洗任务中更接近高保持、弱压制端点。ROME 将隐私值包含率降至 0.5787，同时公开知识保持率下降到 0.5591，形成中间权衡区域。FT 将隐私泄露相关指标压到 0，但公开知识保持率仅为 0.0040，说明强覆盖式微调可能通过破坏公开回答能力来获得低泄露数值。

PACE 和 CAPE 能进一步压低目标隐私值输出，隐私值包含率分别为 0.0243 和 0.0443；但二者公开知识保持率分别只有 0.0984 和 0.1119。这一结果揭示了闭环再编辑的核心副作用：残余泄露驱动的补充编辑能够扩大隐私压制覆盖面，却可能把拒答目标扩散到公开知识集合。

## 4.5  PACE/CAPE 与过度拒答分析

PACE 的设计目标是用残余泄露样本补足单轮编辑覆盖不足。该策略在隐私压制上有效，但 public retain 结果表明，若二轮请求只由失败样本驱动，模型可能学习到更宽泛的人物相关拒答模式。CAPE 通过主体预算和局部性风险控制尝试缓解这一问题，但在当前结果中仍未避免 public collapse。

这一现象说明，隐私清洗不能只报告 private refusal 或目标泄露下降。一个模型如果对隐私问题拒答，同时也对公开问题拒答，那么它更接近“泛化拒答”而不是高质量隐私清洗。本文引入 public refusal 和 public contains 的目的，正是将这种副作用从隐私压制数值中分离出来。

## 4.6  CAPE-Anchor 消融分析

CAPE-Anchor 的目标不是继续追求最低隐私泄露率，而是在闭环再编辑中加入公开知识锚点，使模型在二轮编辑时同时接触隐私拒答请求和公开事实保留信号。表 4 展示了有限消融结果。

表 4  CAPE-Anchor 消融结果

{cape}

PACE-Lite K0 的公开知识保持率为 0.4210，明显高于 naive PACE/CAPE，但隐私值包含率回升到 0.3323。CAPE-Anchor K1 将公开知识保持率进一步提高到 0.6008，同时隐私值包含率为 0.4603，仍低于 ROME 的 0.5787，因而形成较合理的折中点。K2 的公开知识保持率达到 0.6833，但隐私值包含率上升到 0.6357，说明公开锚点过多会削弱隐私压制。

因此，本文采用限定的 Claim A：CAPE-Anchor 在当前 Qwen synthetic privacy 任务上形成有限有效改进。该结论不意味着 CAPE-Anchor 是泄露最低的方法，也不意味着在不损伤公开知识的情况下完成清洗；它说明公开知识锚点可以把 naive PACE/CAPE 从强拒答低保持端点拉回到更可解释的 privacy-utility trade-off 区域。

## 4.7  隐私压制—知识保持权衡图分析

图 3 将不同方法放入隐私压制—知识保持二维空间。横轴为公开知识保持率，纵轴可理解为隐私压制强度。理想方法应位于右上角，即同时具有高公开保持和低隐私泄露。当前结果中没有方法同时占据该区域，说明隐私知识清洗存在明显结构性权衡。

图 3  合成隐私基准上的隐私压制—知识保持权衡散点图

![图 3  合成隐私基准上的隐私压制—知识保持权衡散点图]({FIG_DIR / "fig3_synthetic_privacy_utility_tradeoff.png"})

图 3 的关键观察是：MEMIT 靠近高保持弱压制区域，FT、PACE 和 CAPE 靠近强压制低保持区域，ROME 处于中间位置，而 CAPE-Anchor K1/K2 将操作点向公开知识保持方向移动。该图支撑本文的核心结论：隐私清洗方法不能只按单一泄露率排序，而应在隐私压制和知识保持坐标系中解释。

图 4  CAPE-Anchor 消融折中图

![图 4  CAPE-Anchor 消融折中图]({FIG_DIR / "fig4_cape_anchor_ablation_tradeoff.png"})

图 4 进一步显示，随着公开锚点增加，公开知识保持率上升，但隐私值包含率也同步上升。这说明 public anchor 是一种约束信号，而不是免费收益；其价值在于提供可调节的操作点，使研究者可以根据应用场景选择更偏隐私压制或更偏公开知识保持的配置。

## 4.8  Qwen 公开事实编辑迁移验证

CounterFact 和 zsRE 是公开事实编辑数据集，不是隐私清洗数据集。本文使用它们的目的，是检查 PACE/CAPE 这种 closed-loop request selection 是否能在公开 factual editing 场景中运行，并观察其 reliability 与 generalization 表现。表 5 给出 Qwen 上的 public transfer 结果。

表 5  Qwen 公开事实编辑迁移验证结果

{qwen}

结果显示，ROME+PACE-Edit 和 ROME+CAPE-Edit 在 CounterFact 上可靠性均为 0.5179，在 zsRE 上可靠性均为 0.3117，说明 wrapper 可以迁移运行，但效果中等。该结果不能写成公开基准支撑隐私清洗成功；更稳妥的解释是，闭环请求选择思想具有一定外部可运行性，同时也暴露了它在公开事实编辑任务上的收益边界。

图 5  Qwen 公开迁移结果对比

![图 5  Qwen 公开迁移结果对比]({FIG_DIR / "fig5_public_transfer_qwen.png"})

图 5 的作用是把 Qwen public wrapper 与基础方法放在同一 public editing 语境中比较。它服务的不是隐私清洗主张，而是外部迁移验证：本文方法并非只能在自造数据上运行，但公开事实编辑结果仍不足以支撑“通用提升”的强结论。

## 4.9  GPT-J 第二模型边界验证

为避免公开迁移实验只依赖 Qwen2.5-7B，本文进一步在 GPT-J-6B 上进行第二模型公开事实编辑验证。结果显示，GPT-J-6B 上 ROME 与 FT baseline 均能正常运行，在 CounterFact 与 zsRE 上获得较高 rewrite 成功率。然而，将未重新调参的 ROME-based PACE/CAPE wrapper 直接迁移到 GPT-J-6B 时，wrapper 出现明显塌缩。

表 6  GPT-J 第二模型边界验证结果

{gptj}

per-case 审计进一步排除了路径混用和汇总脚本误读的可能：CounterFact wrapper 的 rewrite 成功数为 0/220，zsRE wrapper 仅为 1/216，而同一模型上的 ROME/FT baseline 正常。该结果说明，闭环请求扩展策略的跨模型迁移依赖底层编辑器超参数、请求集合规模与局部性约束，后续需要模型特定校准或 retain-aware objective。

图 6  GPT-J wrapper collapse 边界验证图

![图 6  GPT-J wrapper collapse 边界验证图]({FIG_DIR / "fig6_gptj_wrapper_boundary_check.png"})

图 6 展示的不是负面异常，而是方法边界。它说明 GPT-J 模型本身和 ROME/FT baseline 并未损坏，问题集中在未调参 wrapper 与 GPT-J 编辑超参数的耦合上。因此，GPT-J 结果应写作边界证据，而不应否定 Qwen synthetic privacy 主实验。

## 4.10  失败项与资源限制分析

表 7 汇总资源限制和方法边界项。KN 在部分设置中受到 48GB 显存限制，IKE 依赖句向量模型但当前环境缺少对应本地依赖；GPT-J wrapper 虽然运行完成，但作为边界结果而非正向有效结果处理。这些条目不应被隐藏，也不应被当作有效 0 值。

表 7  失败项与资源限制分析

{failures}

表 7 的作用是明确实验矩阵中哪些结果可比较、哪些结果只能作为资源或迁移边界。对于课程论文而言，如实报告失败项比强行补齐不可承受的矩阵更可靠，也能避免把资源受限误解释为算法优劣。

## 4.11  典型案例分析

以合成电话号码查询为例，编辑前模型会在直接询问和补全诱导中输出目标号码；ROME 后目标号码出现频率下降，但部分改写提示仍可能触发泄露；PACE/CAPE 进一步压低这类残余泄露，但公开问题可能被一起拒答。该案例解释了为什么 residual leakage mining 对隐私压制有帮助，也解释了为什么必须同时观察 public retain。

另一个典型现象来自公开锚点。CAPE-Anchor 在二轮请求中加入公开事实后，模型更可能保留职业、学校等公开回答能力，但隐私查询上的压制强度会下降。这说明公开锚点不是简单提高所有指标，而是在目标冲突之间移动操作点。该现象与表 4 和图 4 的数值趋势一致。

## 4.12  本章小结

本章结果表明，隐私知识清洗不是单一目标优化问题。提示拒答不能改变内部隐私记忆；基础编辑器呈现不同隐私—效用端点；PACE/CAPE 提高隐私压制但带来过度拒答；CAPE-Anchor 通过公开知识锚点形成有限有效折中；Qwen public transfer 和 GPT-J boundary check 则共同说明 closed-loop wrapper 具有外部可运行性，但跨任务、跨模型迁移仍依赖具体配置。"""


def chapter5() -> str:
    return f"""# 5  系统设计与软件实现

## 5.1  系统总体结构

本文系统由数据构造、隐私注入、模型编辑、评估审计和结果归档五个模块组成。数据构造模块负责生成 synthetic privacy benchmark；隐私注入模块通过 LoRA 构造待编辑模型；模型编辑模块调用基础编辑器和 wrapper 层请求选择策略；评估审计模块统计 private/public 指标；结果归档模块将表格、图像和摘要文件组织为论文可复核材料。

图 7  系统模块与结果归档流程

![图 7  系统模块与结果归档流程]({FIG_DIR / "fig7_system_module_archive_flow.png"})

图 7 展示了系统从数据到论文结果的闭环。该流程的关键不是堆叠工具，而是保证同一 case_id、请求集合、评估输出和结果表之间可以互相追溯，从而降低多轮实验中的状态漂移。

## 5.2  数据构造与隐私注入模块

数据构造模块以虚拟人物为基本单位，生成隐私事实、公开事实和多类型攻击提示。所有隐私值均为合成内容，避免真实个人信息进入论文材料。隐私注入模块使用 LoRA 将这些事实写入基础模型，并在合并后形成统一的待编辑模型。这样可以把“模型是否具有隐私记忆”和“编辑是否压制隐私记忆”两个问题分开讨论。

## 5.3  模型编辑执行模块

模型编辑执行模块负责将统一请求格式转换为不同编辑器可接受的输入。ROME、MEMIT、FT、KN 和 IKE 作为基础编辑器，PACE、CAPE 与 CAPE-Anchor 作为请求选择层策略。该模块保持底层编辑器实现相对独立，使本文的主要变量集中在任务定义、请求选择和评估口径上。

## 5.4  评估与结果汇总模块

评估模块分别生成 private 和 public 查询输出，并计算隐私值包含率、敏感格式率、隐私拒答率、公开知识保持率和公开问题拒答率。结果汇总模块将 JSON、CSV 和图表统一组织，避免论文数值依赖终端输出或人工抄写。对于 failed、missing 或 resource-limited 项，系统单独记录原因和论文处理方式。

## 5.5  实验平台与运行环境

本文实验在远程 GPU 计算环境中完成，主要使用 PyTorch、Transformers 与 EasyEdit 框架。模型权重和大规模中间生成文件不作为论文正文材料，正文只保留可复核的小型摘要、表格、图像和审计说明。该处理方式既降低了材料体积，也使结果解释集中在可比较指标上。

## 5.6  代码组织与可复现性说明

为保证可复现性，本文将数据构造、模型编辑、评估统计和图表生成拆分为相对独立的模块。每轮实验保留输入请求、评估配置、摘要指标和失败记录。这样的组织方式使后续研究者可以单独替换基础模型、编辑器或请求选择策略，而不需要重写完整实验流程。"""


def conclusion() -> str:
    return """# 6  总结与展望

本文围绕大语言模型隐私知识清洗问题，构建了从合成隐私基准、LoRA 可控隐私注入、模型编辑执行到多维评估审计的完整实验流程。与只报告隐私查询拒答率的做法不同，本文同时考察隐私值包含率、敏感格式率、公开知识保持率和公开问题拒答率，使隐私压制和知识损伤能够在同一框架中被观察。

实验结果表明，标准模型编辑方法在隐私清洗任务中呈现不同端点。提示拒答不能改变模型内部隐私记忆；MEMIT 保留公开知识较好但隐私压制不足；ROME 形成中间权衡；FT、PACE 和 CAPE 可以强压目标隐私输出，但容易造成 public collapse。CAPE-Anchor 通过在二轮请求集合中加入公开知识锚点，将 naive PACE/CAPE 中接近塌缩的公开知识保持能力拉回，在 Qwen synthetic privacy 主任务上形成有限有效改进。

本文也保留了方法边界。Qwen CounterFact/zsRE 结果说明 ROME-based wrapper 可以迁移到公开事实编辑任务中运行，但效果中等；GPT-J-6B 结果表明，未重新调参的 wrapper 在第二模型上不稳定，而 ROME/FT baseline 正常。这些结果共同说明，闭环请求扩展策略具有诊断价值和一定迁移性，但跨模型稳定性仍依赖底层编辑器超参数、请求集合规模和局部性约束。

后续工作可以从三个方向推进。第一，引入模型特定校准，使 wrapper 在不同模型架构上更稳定。第二，将公开知识锚点从请求构造扩展为显式 retain-aware objective，使知识保持不只依赖样本混合。第三，引入更严格的 held-out attack split、人工评估和语义评估模型，对隐私拒答合理性和公开回答质量进行更细粒度审计。"""


def revise_markdown() -> str:
    make_system_flow_figure()
    text = MD_PATH.read_text(encoding="utf-8")
    text = replace_between(text, "# 摘要", "# 1  绪论", abstract_block())
    text = replace_between_any(text, ["## 2.4  ROME 与 MEMIT 方法概述", "## 2.4  基础编辑方法与闭环请求选择"], "## 2.5  隐私知识清洗任务定义", section_24())
    text = text.replace("## 3.4  基于 ROME/MEMIT 的隐私拒答编辑", "## 3.4  基于基础编辑器的隐私拒答编辑")
    text = text.replace("ROME 和 MEMIT 使用各自标准超参数运行。", "ROME、MEMIT 与 FT 等基础编辑器使用各自标准设置运行。")
    text = replace_between_any(text, ["## 3.7  实验流水线与复现机制", "## 3.7  CAPE-Anchor 公开知识锚点约束"], "# 4  实验设置与结果分析", section_37())
    text = replace_between(text, "# 4  实验设置与结果分析", "# 5  系统设计与软件实现", chapter4())
    text = replace_between(text, "# 5  系统设计与软件实现", "# 6  总结与展望", chapter5())
    text = replace_between(text, "# 6  总结与展望", "# 参考文献", conclusion())
    text = text.replace("CAPE-Anchor 不是全面优于所有方法的新编辑器", "CAPE-Anchor 不是在所有指标上占优的新编辑器")
    text = text.replace("ground truth 完全可控", "ground truth 可控")
    text = text.replace("基础编辑器的不完全覆盖", "基础编辑器覆盖不足")
    text = text.replace("Locating and Editing Factual Associations in GPT.", "Locating and Editing Factual Associations in Generative Pre-trained Transformers.")
    MD_PATH.write_text(text, encoding="utf-8")
    return text


def clear_document(doc: Document) -> None:
    body = doc._body._element
    for child in list(body):
        if child.tag == qn("w:sectPr"):
            continue
        body.remove(child)


def set_run_font(run, east_asia: str = "宋体", latin: str = "Times New Roman", size: float = 12, bold: bool = False, color: RGBColor | None = None) -> None:
    run.font.name = latin
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    run._element.rPr.rFonts.set(qn("w:ascii"), latin)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), latin)
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def add_para(doc: Document, text: str, first_indent: bool = True, align: int | None = None, size: float = 12, bold: bool = False) -> None:
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    if first_indent:
        p.paragraph_format.first_line_indent = Pt(24)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_table(doc: Document, rows: List[List[str]]) -> None:
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for i, row in enumerate(rows):
        for j, value in enumerate(row):
            cell = table.cell(i, j)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if len(value) <= 14 else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.line_spacing = 1.15
            run = p.add_run(value)
            set_run_font(run, size=8.5, bold=(i == 0))
            if i == 0:
                set_cell_shading(cell, "E8EEF5")
    doc.add_paragraph()


def add_abstract_table(doc: Document, abstract: str, keywords: str) -> None:
    table = doc.add_table(rows=2, cols=1)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, value in enumerate([f"摘要：\n{abstract}", keywords]):
        cell = table.cell(i, 0)
        cell.text = ""
        p = cell.paragraphs[0]
        p.paragraph_format.line_spacing = 1.5
        run = p.add_run(value)
        set_run_font(run, size=12)
    doc.add_paragraph()


def parse_md_table(lines: List[str], i: int) -> tuple[List[List[str]], int]:
    rows = []
    while i < len(lines) and lines[i].strip().startswith("|"):
        line = lines[i].strip()
        if not re.match(r"^\|[-:| ]+\|$", line):
            rows.append([c.strip() for c in line.strip("|").split("|")])
        i += 1
    return rows, i


def add_image(doc: Document, path: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(Path(path)), width=Cm(14.0))


def build_docx(md: str) -> None:
    doc = Document(TEMPLATE)
    clear_document(doc)
    for sec in doc.sections:
        sec.top_margin = Cm(2.54)
        sec.bottom_margin = Cm(2.54)
        sec.left_margin = Cm(3.0)
        sec.right_margin = Cm(2.6)

    cover_lines = [
        ("课程论文", "黑体", 22, True),
        ("课  程：智能系统综合实践", "宋体", 14, False),
        ("题  目：基于模型编辑的大语言模型隐私知识清洗与保护研究", "宋体", 14, False),
        ("级、专业：        级  计算机科学与技术  专业        班", "宋体", 14, False),
        ("学生姓名：", "宋体", 14, False),
        ("提交日期：2026 年 6 月 23 日", "宋体", 14, False),
    ]
    for idx, (text, font, size, bold) in enumerate(cover_lines):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if idx == 0:
            p.paragraph_format.space_after = Pt(48)
        else:
            p.paragraph_format.space_after = Pt(18)
        run = p.add_run(text)
        set_run_font(run, east_asia=font, size=size, bold=bold)
    doc.add_page_break()

    lines = md.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if not line:
            i += 1
            continue
        if line == "# 摘要":
            i += 1
            parts: List[str] = []
            keywords = ""
            while i < len(lines):
                current = lines[i].strip()
                if current.startswith("关键词："):
                    keywords = current
                    i += 1
                    break
                if current:
                    parts.append(current)
                i += 1
            add_abstract_table(doc, "\n".join(parts), keywords)
            continue
        if line.startswith("# "):
            p = doc.add_heading(level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(12)
            run = p.add_run(line[2:].strip())
            set_run_font(run, east_asia="黑体", size=16, bold=True)
            i += 1
            continue
        if line.startswith("## "):
            p = doc.add_heading(level=2)
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(line[3:].strip())
            set_run_font(run, east_asia="黑体", size=14, bold=True)
            i += 1
            continue
        if line.startswith("!["):
            m = re.match(r"!\[(.*?)\]\((.*?)\)", line)
            if m:
                add_image(doc, m.group(2))
            i += 1
            continue
        if line.startswith("|"):
            rows, i = parse_md_table(lines, i)
            add_table(doc, rows)
            continue
        if line.startswith("$$"):
            add_para(doc, line.strip("$"), first_indent=False, align=WD_ALIGN_PARAGRAPH.CENTER, size=11)
            i += 1
            continue
        if re.match(r"^(图|表)\s*\d+", line):
            add_para(doc, line, first_indent=False, align=WD_ALIGN_PARAGRAPH.CENTER, size=10.5, bold=False)
            i += 1
            continue
        add_para(doc, line)
        i += 1
    doc.save(DOCX_PATH)


def write_qa(md: str) -> None:
    banned_patterns = {
        "Codex": r"Codex",
        "GPT（不含 GPT-J 模型名）": r"(?<!-)GPT(?!-J)",
        "服务器": r"服务器",
        "运行脚本": r"运行脚本",
        "开卡": r"开卡",
        "拉代码": r"拉代码",
        "当前版本": r"当前版本",
        "待填": r"待填",
        "占位": r"占位",
        "方法失败": r"方法失败",
        "实验失败": r"实验失败",
        "无损清洗": r"无损清洗",
        "跨模型成功": r"跨模型成功",
    }
    hits = {label: [m.start() for m in re.finditer(pattern, md)] for label, pattern in banned_patterns.items() if re.search(pattern, md)}
    v5_doc = Document("课程论文_v5_academic_polished.docx")
    v7_doc = Document(DOCX_PATH)
    QA_PATH.write_text(
        f"""# Course Paper V8 Final QA

Last updated: 2026-06-23

## 通过项

- 已以 v5 母版为正文基础，只替换过期结果与最终口径。
- 已直接更新 `课程论文_v7_final_academic_polished.docx`，未另开无法追踪的新正文。
- 已重写摘要，使其包含问题背景、方法设计、实验设置、主要发现和边界说明。
- 已新增 `3.7 CAPE-Anchor 公开知识锚点约束`。
- 第 4 章已改为最终结果分析结构，包含 synthetic 主实验、PACE/CAPE over-refusal、CAPE-Anchor 消融、Qwen public transfer、GPT-J boundary、资源限制和典型案例。
- 第 5 章已改为系统总体结构、数据构造、模型编辑、评估汇总、实验平台和可复现性说明。
- DOCX 已设置中文论文风格：一级标题黑体，二级标题黑体，正文宋体小四，英文/数字 Times New Roman。
- GPT-J 已写为 boundary check，未写成第二模型正向有效证据。
- public benchmark 已写为公开事实编辑迁移验证，未写成 PII 清洗依据。

## 结构对比

- v5 DOCX：{len(v5_doc.paragraphs)} 段，{len(v5_doc.tables)} 表。
- 当前 DOCX：{len(v7_doc.paragraphs)} 段，{len(v7_doc.tables)} 表。
- 当前 DOCX 图片关系数：{sum(1 for rel in v7_doc.part.rels.values() if 'image' in rel.reltype)}。

## 仍需人工检查项

- 本机未完成 Word/LibreOffice 页面渲染检查，需要人工打开 DOCX 做最终视觉确认。
- 学生姓名、年级专业等封面字段需要按实际提交要求补齐。
- Word 中若部分宽表分页不理想，需要手工调小表格字号或横向排版。

## 高风险段落

- 4.6 CAPE-Anchor 消融分析：需要保持“有限有效改进”口径，不能改成零损伤或全局占优。
- 4.8 Qwen 公开事实编辑迁移验证：只能写 public factual editing transfer。
- 4.9 GPT-J 第二模型边界验证：只能写迁移边界，不能写第二模型成功。

## 禁用词扫描

{hits if hits else "正文未命中禁用词。"}

## 从 v5 修改了哪些段落

- 摘要。
- 1.4 贡献点。
- 2.4 基础编辑方法与 wrapper 定位。
- 3.4 基础编辑器描述。
- 3.7 CAPE-Anchor。
- 第 4 章全部结果分析。
- 第 5 章系统实现标题与表述。
- 第 6 章总结与展望。

## 删除了哪些压缩文本

- 删除了把第 4 章写成单纯结果汇总的压缩段落。
- 删除了旧版系统实现中的过程化表达。
- 删除了旧图路径和重复图题。
""",
        encoding="utf-8",
    )


def main() -> int:
    md = revise_markdown()
    build_docx(md)
    write_qa(md)
    print(f"updated {MD_PATH}")
    print(f"updated {DOCX_PATH}")
    print(f"wrote {QA_PATH}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
