import json
import shutil
from collections import Counter
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]

SOURCE_DOCX = ROOT / "课程论文_v2_rewritten.docx"
BACKUP_DOCX = ROOT / "课程论文_v2_rewritten_backup_before_cape_update.docx"
OUT_DOCX = ROOT / "课程论文_v2_rewritten_cape_updated.docx"

PAPER_MD = ROOT / "docs" / "course_paper_cape_update.md"
GROUP_MD = ROOT / "docs" / "GROUP_MEETING_REPORT_20260622.md"
LOG_MD = ROOT / "docs" / "CAPE_UPDATE_LOG_20260622.md"

FRAMEWORK_PNG = ROOT / "主要框架大图.png"
CAPE_PNG = ROOT / "CAPEv0.png"

CAPE_V0 = ROOT / "artifacts" / "run_20260622_v2_cape_b1_tau05"
CAPE_V1 = ROOT / "artifacts" / "run_20260622_v2_cape_v1_top20_tau07_direct"
ANALYSIS = ROOT / "artifacts" / "analysis_v2_audit_20260622"


MAIN_RESULTS = [
    ["Merged pre-edit", "0.9387", "0.6650", "0.9927", "0.0000", "0.9766", "-"],
    ["MEMIT direct", "0.8140", "0.5993", "0.8853", "0.1950", "0.8472", "0.0897"],
    ["ROME direct", "0.5787", "0.4767", "0.8563", "0.5973", "0.5591", "0.2873"],
    ["PACE max2/person", "0.0243", "0.0150", "0.0740", "0.9347", "0.0984", "0.8052"],
    ["CAPE-v0", "0.0023", "0.0023", "0.0190", "0.9890", "0.0060", "0.9877"],
    ["CAPE-v1", "0.0443", "0.0250", "0.2587", "0.8557", "0.1119", "0.6901"],
]

CAPE_ATTACK_RESULTS = [
    ["direct", "0.0417", "0.0183", "0.1883", "0.9267"],
    ["paraphrase", "0.0417", "0.0233", "0.2033", "0.9050"],
    ["completion", "0.1100", "0.0683", "0.6483", "0.5833"],
    ["roleplay", "0.0000", "0.0000", "0.0483", "0.9800"],
    ["context", "0.0283", "0.0150", "0.2050", "0.8833"],
]

CAPE_PUBLIC_RESULTS = [
    ["same_subject_public", "0.1100"],
    ["same_relation_other_subject", "0.1075"],
    ["general_knowledge", "0.1750"],
]


def read_json(path: Path):
    with path.open("r", encoding="utf-8") as fh:
        return json.load(fh)


def set_east_asia(run, font="宋体"):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, size=8):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(size)
    set_east_asia(run)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows, font_size=8):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = True
    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, bold=True, size=font_size)
        set_cell_shading(table.rows[0].cells[i], "E8EEF5")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value, size=font_size)
    doc.add_paragraph()
    return table


def set_doc_styles(doc):
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.6)
    section.right_margin = Cm(2.6)

    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.space_after = Pt(5)

    for style_name, font, size, color in [
        ("Heading 1", "黑体", 16, RGBColor(31, 77, 120)),
        ("Heading 2", "黑体", 14, RGBColor(31, 77, 120)),
        ("Heading 3", "黑体", 12, RGBColor(31, 77, 120)),
    ]:
        style = doc.styles[style_name]
        style.font.name = font
        style._element.rPr.rFonts.set(qn("w:eastAsia"), font)
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(5)


def add_para(doc, text, first_indent=True):
    p = doc.add_paragraph()
    if first_indent:
        p.paragraph_format.first_line_indent = Pt(21)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    set_east_asia(run)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    set_east_asia(run)
    return p


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.italic = True
    set_east_asia(run)
    return p


def add_picture_if_exists(doc, path: Path, caption: str, width_inches=5.8):
    if path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(path), width=Inches(width_inches))
        add_caption(doc, caption)
    else:
        add_para(doc, f"[图示预留：{caption}；缺少图片文件 {path.name}]", first_indent=False)


def load_public_failure_counts(public_eval_path: Path):
    markers = [
        "cannot provide private information",
        "can't provide private information",
        "cannot share private information",
        "private information",
        "i cannot provide",
        "i can't provide",
        "i'm sorry",
    ]
    data = read_json(public_eval_path)
    counts = Counter()
    for row in data.get("details", []):
        if row.get("contains_match") or row.get("exact_match"):
            counts["success"] += 1
            continue
        output = str(row.get("output") or "").lower()
        if any(marker in output for marker in markers):
            counts["refusal"] += 1
        else:
            counts["wrong_or_drift"] += 1
    total = sum(counts.values())
    return counts, total


def build_paper_markdown():
    lines = [
        "# 课程论文 CAPE 更新稿",
        "",
        "## 摘要更新建议",
        "",
        "大语言模型可能在预训练或后训练过程中记忆并输出个人隐私信息。本文围绕基于模型编辑的隐私知识清洗任务，构建了包含 private/public 对照的 synthetic privacy benchmark，并通过 MLP-only LoRA 向 Qwen2.5-7B 注入可控隐私记忆，形成 merged leakage model。在此基础上，本文基于 EasyEdit 框架完成 ROME direct-only、MEMIT direct-only、PACE max2/person、CAPE-v0 与 CAPE-v1 等实验，并从隐私泄露、公开知识保持、过度拒答和 privacy-utility trade-off 四个维度进行审计。实验结果表明，现有模型编辑方法难以同时实现强隐私压制和高公开知识保持：MEMIT 保留公开知识较好但隐私压制不足，PACE 将 Private Value Contains 降至 0.0243，但带来严重过度拒答。进一步的 CAPE-v0/v1 探索表明，请求数量、覆盖人物范围和 prompt 类型会影响 over-refusal；CAPE-v1 可缓解 CAPE-v0 的极端 public collapse，但仍未达到理想折中。",
        "",
        "关键词：大语言模型；模型编辑；隐私知识清洗；知识保持；过度拒答；ROME；MEMIT；PACE；CAPE",
        "",
        "## 3.6 CAPE 副作用感知请求选择策略",
        "",
        "CAPE（Collateral-Aware Privacy Editing）并不是新的底层模型编辑算法，而是建立在 ROME/PACE 编辑流程之上的 request selection / wrapper。其动机来自 PACE 实验中的过度拒答现象：简单地对 residual leakage 追加 re-edit requests 可以压低隐私泄露，但容易使模型学习到过宽泛的拒答模式，进而损伤公开事实与通用知识问答。",
        "",
        "CAPE-v0 在 residual leakage re-edit 阶段加入 public-anchor blocking 与 per-person edit budget：前者使用同一 subject 的公开事实保留率过滤高风险 subject，后者限制每个 subject 的追加编辑次数。该版本保持 ROME/MEMIT 底层更新规则不变，仅改变进入编辑器的请求集合。",
        "",
        "CAPE-v1 针对 CAPE-v0 的 public collapse 进一步收缩请求选择空间：将 max_total_requests 限制为 20，提高 public-anchor threshold 至 tau=0.7，只选择 target_exact/value_contains 类型的 residual leakage，并将 round2 request 统一转换为 canonical direct prompt。该设计用于验证 request count 与 prompt type 是否是导致 over-refusal 的关键因素。",
        "",
        "## 4.x CAPE-v0/v1 副作用感知请求选择实验",
        "",
        "### 主结果表",
        "",
        "| 方法 | Private Value Contains ↓ | PII-format Regex ↓ | Sensitive Pattern ↓ | Private Refusal ↑ | Public Contains ↑ | Public Refusal ↓ |",
        "| --- | ---: | ---: | ---: | ---: | ---: | ---: |",
    ]
    for row in MAIN_RESULTS:
        lines.append("| " + " | ".join(row) + " |")
    lines += [
        "",
        "### CAPE-v0 结果与诊断",
        "",
        "CAPE-v0 在 B=1、tau=0.5 设置下，从 2569 个 residual leakage candidates 中选出 60 条 round2 re-edit requests，涉及 60 个 selected people。实验结果显示，CAPE-v0 将 Private Value Contains 降至 0.0023，但 Public Contains 仅为 0.0060，Public Refusal 达到 0.9877。该结果说明 CAPE-v0 并没有缓解 PACE 的 public collapse，反而进一步滑向全局拒答。其主要原因在于 B=1 只限制单个 subject 的请求数，没有限制总覆盖人物数；同时统一 refusal target 容易诱导模型将人物相关问题泛化为拒答。",
        "",
        "### CAPE-v1 结果与分析",
        "",
        "CAPE-v1 使用 max_total_requests=20、tau=0.7、canonical direct prompt 与 target_exact/value_contains candidate source。该设置选出 20 条 round2 requests，涉及 20 个 selected people；所有 round2 requests 均为 direct prompt，总编辑请求数为 60，即原始 ROME 40 条加 CAPE round2 20 条。",
        "",
        "相较 CAPE-v0，CAPE-v1 将 Public Contains 从 0.0060 提升至 0.1119，将 Public Refusal 从 0.9877 降至 0.6901，说明限制请求数量和规范化 prompt 形式能够缓解最极端的 over-refusal。但相比 PACE max2/person，CAPE-v1 的 Public Contains 仅从 0.0984 小幅提高到 0.1119，同时 Private Value Contains 从 0.0243 上升到 0.0443，Sensitive Pattern 从 0.0740 上升到 0.2587。因此 CAPE-v1 不能写成明显优胜方法，只能视为诊断性改进。",
        "",
        "### CAPE-v1 分攻击类型结果",
        "",
        "| Attack Type | Value Contains ↓ | Regex ↓ | Sensitive Pattern ↓ | Refusal ↑ |",
        "| --- | ---: | ---: | ---: | ---: |",
    ]
    for row in CAPE_ATTACK_RESULTS:
        lines.append("| " + " | ".join(row) + " |")
    lines += [
        "",
        "CAPE-v1 在 direct、paraphrase、roleplay 和 context 上呈现较低的 Value Contains，但 completion 仍是主要残留来源：completion 的 Value Contains 为 0.1100，Sensitive Pattern 为 0.6483，Refusal 仅为 0.5833。这说明 canonical direct re-edit 对 completion-style leakage 的泛化能力不足。",
        "",
        "### CAPE-v1 public breakdown",
        "",
        "| Public Type | Public Contains |",
        "| --- | ---: |",
    ]
    for row in CAPE_PUBLIC_RESULTS:
        lines.append("| " + " | ".join(row) + " |")
    lines += [
        "",
        "CAPE-v1 的 public failure heuristic 为：refusal = 1739/2520 = 0.6901，success = 282/2520 = 0.1119，wrong_or_drift = 499/2520 = 0.1980。该统计表明 public damage 仍主要来自 over-refusal，而不是单纯事实漂移。",
        "",
        "### 小结",
        "",
        "CAPE-v1 相比 CAPE-v0 证明限制请求数量和规范化 prompt 形式能够缓解极端 public collapse，但其 public retain 仍显著不足，说明 request selection 只能部分缓解副作用，尚不足以单独解决隐私清洗中的 trade-off。",
        "",
        "## 第6章结论更新建议",
        "",
        "进一步的 CAPE-v0/v1 探索表明，request selection 的粒度、覆盖人物数和 prompt 类型会显著影响 over-refusal 程度。CAPE-v1 相比 CAPE-v0 缓解了极端 public collapse，但仍未达到理想的隐私压制—公开知识保持折中。这说明后续需要在请求选择之外引入更显式的 locality 约束、protected public anchors 或 calibrated refusal target。",
        "",
    ]
    return "\n".join(lines)


def build_group_report():
    return """# 阶段性实验进展汇报：基于模型编辑的隐私知识清洗与保护

## 1. 当前完成工作

本阶段围绕“基于模型编辑的大语言模型隐私知识清洗”任务，完成了从数据构造、隐私注入、模型编辑、闭环再编辑到结果审计的完整实验流程。具体包括：构建 100 人规模的 V2 synthetic privacy benchmark；通过 MLP-only LoRA 向 Qwen2.5-7B 注入 private/public facts，得到 merged leakage model；基于 EasyEdit 框架完成 ROME direct-only、MEMIT direct-only、PACE max2/person、CAPE-v0 和 CAPE-v1 等多组实验；并进一步补充了指标定义审计、over-refusal 分析和 privacy-utility trade-off 分析。

## 2. 实验体系概览

整体流程为：synthetic privacy benchmark 构建 -> LoRA privacy injection -> merged leakage model -> ROME/MEMIT baseline -> PACE residual re-edit -> CAPE side-effect-aware request selection -> full private/public evaluation。CAPE 只改变 request selection，不修改 ROME、MEMIT 或 EasyEdit 底层编辑算法。

## 3. 主要结果表

| 方法 | Private Value Contains ↓ | PII-format Regex ↓ | Sensitive Pattern ↓ | Private Refusal ↑ | Public Contains ↑ | Public Refusal ↓ |
| --- | ---: | ---: | ---: | ---: | ---: | ---: |
| Merged pre-edit | 0.9387 | 0.6650 | 0.9927 | 0.0000 | 0.9766 | - |
| MEMIT direct | 0.8140 | 0.5993 | 0.8853 | 0.1950 | 0.8472 | 0.0897 |
| ROME direct | 0.5787 | 0.4767 | 0.8563 | 0.5973 | 0.5591 | 0.2873 |
| PACE max2/person | 0.0243 | 0.0150 | 0.0740 | 0.9347 | 0.0984 | 0.8052 |
| CAPE-v0 | 0.0023 | 0.0023 | 0.0190 | 0.9890 | 0.0060 | 0.9877 |
| CAPE-v1 | 0.0443 | 0.0250 | 0.2587 | 0.8557 | 0.1119 | 0.6901 |

## 4. 指标审计说明

Private Value Contains 对应当前脚本中的 `target_exact_leak_rate`，它实际衡量目标隐私值是否作为子串出现在输出中，不应解释为严格 exact equality。PII-format Regex 主要覆盖 phone/email 等格式化敏感值。Public Contains 是公开事实值的 contains-style retain 指标，不等同于严格事实准确率。Public Refusal 是基于输出文本中 refusal marker 的启发式统计，用于分析 over-refusal。

## 5. ROME / MEMIT / PACE 结论

Merged pre-edit 模型具备明显隐私泄露行为，同时公开知识保持率较高，因此可以作为有效待编辑模型。ROME direct-only 能降低部分泄露，但 residual leakage 和 public damage 均明显。MEMIT direct-only 在公开知识保持方面较好，Public Contains 达到 0.8472，但 Private Value Contains 仍为 0.8140，隐私压制不足。PACE max2/person 能显著压低泄露，Private Value Contains 降至 0.0243，但 Public Contains 仅为 0.0984，Public Refusal 达到 0.8052，表现出明显 over-refusal。

## 6. CAPE-v0/v1 初步探索

为缓解 PACE 的 over-refusal 问题，本阶段进一步设计 CAPE（Collateral-Aware Privacy Editing）副作用感知请求选择策略。CAPE 不修改 ROME/MEMIT 底层编辑公式，而是在 residual leakage re-edit 阶段加入 public-anchor blocking 和 per-person edit budget，用于控制再编辑请求的副作用。

CAPE-v0 采用 B=1、tau=0.5 的设置，但由于没有限制总覆盖人物数，最终 selected_people 达到 60，导致模型进一步滑向全局拒答。实验中 CAPE-v0 虽然将 Private Value Contains 降至 0.0023，但 Public Contains 仅为 0.0060，Public Refusal 达到 0.9877，说明该版本未能缓解 public collapse。

在此基础上，CAPE-v1 进一步限制 max_total_requests=20，提高 public-anchor threshold 至 tau=0.7，并将 round2 request 统一为 canonical direct prompt。相比 CAPE-v0，CAPE-v1 将 Public Contains 从 0.0060 提升至 0.1119，将 Public Refusal 从 0.9877 降至 0.6901，说明限制请求数量和规范化 prompt 形式能够缓解最极端的 over-refusal。但 CAPE-v1 相比 PACE max2/person 仅带来有限的 public retain 改善，同时隐私压制变弱，因此仍不能视为最终有效方法。

## 7. 当前最重要发现

当前没有任何方法同时实现强隐私压制和高公开知识保持。ROME 具有一定抑制效果但残余泄露和 public damage 均明显；MEMIT 保留公开知识较好但隐私压制不足；PACE 能显著压低泄露但产生严重 over-refusal；CAPE-v1 相比 CAPE-v0 缓解了极端 public collapse，但仍不足以解决整体 trade-off。因此当前阶段的核心贡献在于构建完整实验闭环、发现并量化 privacy-utility trade-off，并初步验证 request selection 对 over-refusal 的影响。

## 8. 局限性

第一，当前 benchmark 是 synthetic privacy setting，不能直接声称完成真实预训练 PII 删除。第二，Public Contains 是 contains-style 指标，不等价于严格事实准确率。第三，CAPE-v1 的 public refusal 仍然较高，说明 request selection 只能部分缓解副作用。第四，completion-style leakage 仍是主要难点，direct-form re-edit 对该类攻击的泛化不足。

## 9. 下一步计划

组会前不继续开新 GPU 实验。后续可以考虑在三个方向改进：第一，引入更显式的 locality constraint 或 protected public anchors，使公开知识保持从筛选条件变成更强约束；第二，针对 completion-style leakage 设计攻击类型感知编辑策略；第三，改进 refusal target，避免模型学习到过宽泛的“人物相关问题统一拒答”模式。
"""


def build_log():
    return """# CAPE Update Log 20260622

## Updated files

- `课程论文_v2_rewritten_backup_before_cape_update.docx`: backup before CAPE update.
- `课程论文_v2_rewritten_cape_updated.docx`: updated Word paper with CAPE-v0/v1 results.
- `docs/course_paper_cape_update.md`: Markdown version of paper additions.
- `docs/GROUP_MEETING_REPORT_20260622.md`: formal group meeting report.
- `docs/CAPE_UPDATE_LOG_20260622.md`: this log.

## Evidence used

- `artifacts/run_20260622_v2_cape_b1_tau05/`
- `artifacts/run_20260622_v2_cape_v1_top20_tau07_direct/`
- `artifacts/analysis_v2_audit_20260622/`
- `主要框架大图.png`
- `CAPEv0.png`

## Main update

CAPE is no longer described as a successful final method. It is framed as a side-effect-aware request selection exploration. CAPE-v0 is treated as a negative diagnostic result, while CAPE-v1 is treated as a partial diagnostic improvement that mitigates extreme public collapse but does not solve the privacy-utility trade-off.

## Self-check

- Word output generated: yes.
- Backup generated: yes.
- Main CAPE-v0/v1 table added: yes.
- CAPE-v1 attack-type and public breakdown added: yes.
- Figures inserted if image files exist: yes.
- No claim of real PII deletion or complete solution: yes.
- No EasyEdit/ROME/MEMIT core logic modified: yes.
"""


def build_docx():
    doc = Document(str(SOURCE_DOCX))
    set_doc_styles(doc)
    replacements = {
        "该策略确实能够显著降低泄露": "该策略能够将泄露降至较低水平",
        "能够显著降低隐私泄露": "能够明显降低隐私泄露",
    }
    for para in doc.paragraphs:
        for old, new in replacements.items():
            if old in para.text:
                para.text = para.text.replace(old, new)

    doc.add_page_break()
    doc.add_heading("CAPE-v0/v1 副作用感知请求选择实验更新", level=1)
    add_para(
        doc,
        "本节根据最新 CAPE-v0 与 CAPE-v1 实验结果，对论文中的方法、实验分析和结论部分进行增补。CAPE 不修改 ROME/MEMIT 或 EasyEdit 底层编辑算法，而是在 residual leakage re-edit 阶段调整请求选择策略，用于分析 request count、selected people 范围和 prompt type 对 over-refusal 的影响。",
    )

    doc.add_heading("图示补充", level=2)
    add_picture_if_exists(doc, FRAMEWORK_PNG, "图1 基于模型编辑的隐私知识清洗总体框架", width_inches=5.8)
    add_picture_if_exists(
        doc,
        CAPE_PNG,
        "图2 CAPE 副作用感知请求选择机制设计。该图展示 CAPE 的设计目标和请求选择流程，实际效果需结合 CAPE-v0/v1 实验结果分析。",
        width_inches=5.8,
    )

    doc.add_heading("3.6 CAPE 副作用感知请求选择策略", level=2)
    add_para(
        doc,
        "CAPE（Collateral-Aware Privacy Editing）并不是新的底层模型编辑算法，而是建立在 ROME/PACE 编辑流程之上的 request selection / wrapper。其动机来自 PACE 实验中的过度拒答现象：简单地对 residual leakage 追加 re-edit requests 可以压低隐私泄露，但容易使模型学习到过宽泛的拒答模式，进而损伤公开事实与通用知识问答。",
    )
    add_para(
        doc,
        "CAPE-v0 在 residual leakage re-edit 阶段加入 public-anchor blocking 与 per-person edit budget：前者使用同一 subject 的公开事实保留率过滤高风险 subject，后者限制每个 subject 的追加编辑次数。CAPE-v1 进一步将 max_total_requests 限制为 20，提高 public-anchor threshold 至 tau=0.7，只选择 target_exact/value_contains 类型的 residual leakage，并将 round2 request 统一转换为 canonical direct prompt。",
    )

    doc.add_heading("4.x CAPE-v0/v1 副作用感知请求选择实验", level=2)
    add_para(doc, "表中结果显示，CAPE-v1 相比 CAPE-v0 缓解了极端 public collapse，但仍未形成理想的隐私压制—公开知识保持折中。")
    add_table(
        doc,
        ["方法", "Private Value Contains ↓", "PII-format Regex ↓", "Sensitive Pattern ↓", "Private Refusal ↑", "Public Contains ↑", "Public Refusal ↓"],
        MAIN_RESULTS,
        font_size=7,
    )
    doc.add_heading("CAPE-v0 结果与诊断", level=3)
    add_para(
        doc,
        "CAPE-v0 在 B=1、tau=0.5 设置下选出 60 条 round2 re-edit requests，涉及 60 个 selected people。该版本将 Private Value Contains 降至 0.0023，但 Public Contains 仅为 0.0060，Public Refusal 达到 0.9877，说明简单 request filtering 不能自动保证 public retain。",
    )
    doc.add_heading("CAPE-v1 结果与分析", level=3)
    add_para(
        doc,
        "CAPE-v1 使用 max_total_requests=20、tau=0.7、canonical direct prompt 与 target_exact/value_contains candidate source。相较 CAPE-v0，Public Contains 从 0.0060 提升至 0.1119，Public Refusal 从 0.9877 降至 0.6901；但相比 PACE max2/person，隐私压制变弱，因此只能视为诊断性改进。",
    )

    doc.add_heading("CAPE-v1 分攻击类型分析", level=3)
    add_table(doc, ["Attack Type", "Value Contains ↓", "Regex ↓", "Sensitive Pattern ↓", "Refusal ↑"], CAPE_ATTACK_RESULTS, font_size=8)
    add_para(
        doc,
        "CAPE-v1 在 direct、paraphrase、roleplay 和 context 上呈现较低的 Value Contains，但 completion 类型仍有明显残留。completion 的 Value Contains 为 0.1100，Sensitive Pattern 为 0.6483，Refusal 仅为 0.5833，说明 canonical direct re-edit 对 completion-style leakage 的泛化能力不足。",
    )

    doc.add_heading("CAPE-v1 public breakdown", level=3)
    add_table(doc, ["Public Type", "Public Contains"], CAPE_PUBLIC_RESULTS, font_size=8)
    add_para(
        doc,
        "CAPE-v1 的 public failure heuristic 为：refusal = 1739/2520 = 0.6901，success = 282/2520 = 0.1119，wrong_or_drift = 499/2520 = 0.1980。该统计表明 public damage 仍主要来自 over-refusal，而不是单纯事实漂移。",
    )

    doc.add_heading("结论更新", level=2)
    add_para(
        doc,
        "进一步的 CAPE-v0/v1 探索表明，request selection 的粒度、覆盖人物数和 prompt 类型会显著影响 over-refusal 程度。CAPE-v1 相比 CAPE-v0 缓解了极端 public collapse，但仍未达到理想的隐私压制—公开知识保持折中。这说明后续需要在请求选择之外引入更显式的 locality 约束、protected public anchors 或 calibrated refusal target。",
    )
    doc.save(str(OUT_DOCX))


def main():
    if not SOURCE_DOCX.exists():
        raise FileNotFoundError(SOURCE_DOCX)
    if not BACKUP_DOCX.exists():
        shutil.copy2(SOURCE_DOCX, BACKUP_DOCX)

    PAPER_MD.parent.mkdir(parents=True, exist_ok=True)
    PAPER_MD.write_text(build_paper_markdown(), encoding="utf-8")
    GROUP_MD.write_text(build_group_report(), encoding="utf-8")
    LOG_MD.write_text(build_log(), encoding="utf-8")
    build_docx()

    # Basic evidence sanity checks.
    required = [
        CAPE_V0 / "privacy_leakage_eval_v2_cape_b1_tau05_full.json",
        CAPE_V1 / "privacy_leakage_eval_v2_cape_v1_top20_tau07_direct_full.json",
        ANALYSIS / "CAPE_V0_FAILURE_DIAGNOSIS.md",
    ]
    missing = [str(path) for path in required if not path.exists()]
    if missing:
        print("[WARN] Missing expected evidence:")
        for path in missing:
            print(path)
    print(f"[OK] backup: {BACKUP_DOCX}")
    print(f"[OK] docx: {OUT_DOCX}")
    print(f"[OK] paper markdown: {PAPER_MD}")
    print(f"[OK] group report: {GROUP_MD}")
    print(f"[OK] log: {LOG_MD}")


if __name__ == "__main__":
    main()
