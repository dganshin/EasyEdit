from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn


MD_PATH = Path("docs/course_paper_v6_results_placeholders.md")
DOCX_PATH = Path("课程论文_v6_results_placeholders.docx")


def set_run_font(run, east_asia: str = "宋体", latin: str = "Times New Roman", size: int = 11) -> None:
    run.font.name = latin
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    run.font.size = Pt(size)


def add_paragraph(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    set_run_font(run)


def main() -> int:
    if not MD_PATH.exists():
        raise FileNotFoundError(MD_PATH)
    document = Document()
    styles = document.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    styles["Normal"].font.size = Pt(11)

    for raw in MD_PATH.read_text(encoding="utf-8").splitlines():
        line = raw.strip()
        if not line:
            continue
        if line.startswith("# "):
            document.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            document.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            document.add_heading(line[4:], level=3)
        elif line.startswith("- "):
            paragraph = document.add_paragraph(style="List Bullet")
            run = paragraph.add_run(line[2:])
            set_run_font(run)
        elif line.startswith("> "):
            paragraph = document.add_paragraph()
            run = paragraph.add_run(line[2:])
            run.italic = True
            set_run_font(run)
        else:
            add_paragraph(document, line)

    document.save(DOCX_PATH)
    print(f"docx: {DOCX_PATH}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
