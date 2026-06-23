#!/usr/bin/env python3
"""Build v7 final academic polished course paper.

The script uses frozen result artifacts only. It does not run experiments.
Outputs:
- docs/course_paper_v7_final_academic_polished.md
- 课程论文_v7_final_academic_polished.docx
- docs/COURSE_PAPER_V7_FINAL_CHANGELOG.md
- docs/PAPER_FINAL_QUALITY_CHECKLIST.md
"""

from __future__ import annotations

import csv
import re
from pathlib import Path
from typing import Any, Dict, Iterable, List, Sequence

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(".")
OUT_MD = Path("docs/course_paper_v7_final_academic_polished.md")
OUT_DOCX = Path("课程论文_v7_final_academic_polished.docx")
CHANGELOG = Path("docs/COURSE_PAPER_V7_FINAL_CHANGELOG.md")
CHECKLIST = Path("docs/PAPER_FINAL_QUALITY_CHECKLIST.md")
ASSET_DIR = Path("artifacts/final_paper_assets_20260623")
FIG_DIR = ASSET_DIR / "figures"
TABLE_DIR = ASSET_DIR / "tables"


def read_csv(path: Path) -> List[Dict[str, str]]:
    with path.open(encoding="utf-8", newline="") as fh:
        return list(csv.DictReader(fh))


def md_table(rows: List[Dict[str, Any]], fields: Sequence[str]) -> str:
    lines = ["|" + "|".join(fields) + "|", "|" + "|".join(["---"] * len(fields)) + "|"]
    for row in rows:
        lines.append("|" + "|".join(str(row.get(field, "")) for field in fields) + "|")
    return "\n".join(lines)


def fmt_float(value: str) -> str:
    if value in ("", "missing", "failed"):
        return value
    try:
        return f"{float(value):.4f}"
    except Exception:
        return value


def load_tables() -> Dict[str, List[Dict[str, str]]]:
    return {
        "synthetic": read_csv(TABLE_DIR / "table1_synthetic_privacy_main.csv"),
        "cape": read_csv(TABLE_DIR / "table2_cape_anchor_ablation.csv"),
        "qwen": read_csv(TABLE_DIR / "table3_qwen_public_transfer.csv"),
        "gptj": read_csv(TABLE_DIR / "table4_gptj_boundary_check.csv"),
        "failures": read_csv(TABLE_DIR / "table5_failure_and_resource_limits.csv"),
    }


def compact_synthetic(rows: List[Dict[str, str]]) -> List[Dict[str, str]]:
    keep = [
        "Merged pre-edit",
        "Prompt Refusal",
        "ROME",
        "MEMIT",
        "FT",
        "PACE",
        "CAPE",
        "PACE_LITE_B20_K0",
        "CAPE_ANCHOR_B20_K1",
        "CAPE_ANCHOR_B20_K2",
        "KN",
        "IKE",
    ]
    by_method = {r["Method"]: r for r in rows}
    labels = {
        "Merged pre-edit": "合并后泄露模型",
        "Prompt Refusal": "提示拒答",
        "PACE_LITE_B20_K0": "PACE-Lite",
        "CAPE_ANCHOR_B20_K1": "CAPE-Anchor K1",
        "CAPE_ANCHOR_B20_K2": "CAPE-Anchor K2",
    }
    out = []
    for method in keep:
        r = by_method[method]
        status = r["Status"]
        out.append(
            {
                "方法": labels.get(method, method),
                "状态": status,
                "隐私值包含率↓": r["Private leak ↓"] if status == "ok" else status,
                "PII格式率↓": r["PII regex ↓"] if status == "ok" else status,
                "隐私拒答率": r["Private refusal"] if status == "ok" else status,
                "公开知识保持率↑": r["Public contains ↑"] if status == "ok" else status,
            }
        )
    return out


def compact_cape(rows: List[Dict[str, str]]) -> List[Dict[str, str]]:
    labels = {
        "PACE_LITE_B20_K0": "PACE-Lite K0（无公开锚点）",
        "CAPE_ANCHOR_B20_K1": "CAPE-Anchor K1",
        "CAPE_ANCHOR_B20_K2": "CAPE-Anchor K2",
    }
    return [
        {
            "配置": labels.get(r["Variant"], r["Variant"]),
            "隐私值包含率↓": r["Private leak ↓"],
            "PII格式率↓": r["PII regex ↓"],
            "隐私拒答率": r["Private refusal"],
            "公开知识保持率↑": r["Public contains ↑"],
            "解释": "较均衡折中" if r["Variant"].endswith("K1") else ("公开保持更高但隐私压制更弱" if r["Variant"].endswith("K2") else "无锚点对照"),
        }
        for r in rows
    ]


def compact_qwen(rows: List[Dict[str, str]]) -> List[Dict[str, str]]:
    out = []
    for r in rows:
        if r["Status"].startswith("failed"):
            continue
        out.append(
            {
                "数据集": "CounterFact" if r["Dataset"] == "counterfact" else "zsRE",
                "方法": r["Method"].replace("ROME_PACE_EDIT", "ROME+PACE-Edit").replace("ROME_CAPE_EDIT", "ROME+CAPE-Edit"),
                "样本数": r["Cases"],
                "可靠性↑": r["Reliability ↑"],
                "泛化性↑": r["Generalization ↑"],
            }
        )
    return out


def compact_gptj(rows: List[Dict[str, str]]) -> List[Dict[str, str]]:
    return [
        {
            "数据集": "CounterFact" if r["Dataset"] == "counterfact" else "zsRE",
            "方法": r["Method"].replace("ROME_PACE_EDIT", "ROME+PACE-Edit").replace("ROME_CAPE_EDIT", "ROME+CAPE-Edit"),
            "可靠性↑": r["Reliability ↑"],
            "泛化性↑": r["Generalization ↑"],
            "逐样本rewrite": r["Per-case rewrite"] or "未统计",
            "解释": "基础编辑正常" if r["Interpretation"] == "baseline normal" else "wrapper塌缩",
        }
        for r in rows
    ]


def compact_failures(rows: List[Dict[str, str]]) -> List[Dict[str, str]]:
    return [
        {
            "范围": r["Scope"],
            "方法/现象": r["Method"],
            "状态": r["Status"],
            "原因": r["Reason"],
            "论文处理": r["Policy"],
        }
        for r in rows
    ]


def build_markdown(tables: Dict[str, List[Dict[str, str]]]) -> str:
    synthetic = compact_synthetic(tables["synthetic"])
    cape = compact_cape(tables["cape"])
    qwen = compact_qwen(tables["qwen"])
    gptj = compact_gptj(tables["gptj"])
    failures = compact_failures(tables["failures"])

    return f"""# 基于模型编辑的大语言模型隐私知识清洗与保护研究

课程：智能系统综合实践

类型：研究型课程论文

关键词：大语言模型；模型编辑；隐私知识清洗；知识保持；过度拒答

# 摘要

大语言模型在预训练和后训练过程中可能形成对个人敏感信息的参数化记忆，并在特定提示下输出电话号码、邮箱等个人可识别信息（Personally Identifiable Information，PII）。仅依赖输出过滤或提示层拒答难以改变模型内部已经形成的知识关联，而全量重训又需要较高计算成本和数据访问条件。针对这一问题，本文研究基于模型编辑的隐私知识清洗方法，目标是在降低目标隐私泄露的同时保持同主体公开知识、同关系其他主体公开知识和通用知识。本文首先构建 private/public 解耦的合成隐私评测基准，将隐私事实与公开事实显式分离；随后通过低秩适配（Low-Rank Adaptation，LoRA）向 Qwen2.5-7B 注入可控隐私记忆，形成统一的合并后隐私泄露模型。在此基础上，本文将秩一模型编辑（Rank-One Model Editing，ROME）、Transformer 批量记忆编辑（Mass-Editing Memory in a Transformer，MEMIT）和微调（Fine-Tuning，FT）等方法适配到隐私拒答式清洗任务，并设计隐私感知闭环编辑（Privacy-Aware Closed-loop Editing，PACE）、副作用感知隐私编辑（Collateral-Aware Privacy Editing，CAPE）以及引入公开知识锚点的 CAPE-Anchor。实验结果表明，单纯提示拒答难以改变模型内部隐私记忆；ROME 能降低部分隐私泄露但会影响公开知识；MEMIT 公开知识保持较好但隐私压制较弱；FT 与 PACE/CAPE 可形成强隐私压制端点，但容易造成公开知识保持下降和过度拒答。CAPE-Anchor 通过显式加入公开知识锚点，将 naive PACE/CAPE 中接近塌缩的公开知识保持能力拉回，在牺牲部分隐私压制强度的情况下形成了更合理的隐私压制—知识保持折中。进一步的 Qwen 公开事实编辑实验表明，closed-loop request selection 可以迁移运行但效果中等；GPT-J-6B 上的边界实验显示，未重新调参的 wrapper 对模型和编辑器超参数较敏感。本文结果为模型特定校准、局部性约束和 retain-aware objective 的后续研究提供了可复现实验依据。

# 关键词

大语言模型；模型编辑；隐私知识清洗；ROME；MEMIT；PACE；CAPE-Anchor；过度拒答

# 第1章 绪论

## 1.1 研究背景

大语言模型已经成为智能系统的重要基础设施，被广泛应用于问答、办公自动化、代码生成和知识服务等场景。模型通过大规模语料学习语言规律和事实关联，在获得强生成能力的同时，也可能记忆训练语料或后训练数据中的具体事实。若这些事实涉及电话号码、邮箱、住址、工作单位等敏感信息，模型在特定提示下可能输出隐私内容，进而带来安全、合规和信任风险。

传统防护方式多发生在输出层，例如关键词过滤、安全分类器和拒答模板。这类方法部署便利，但没有直接改变模型内部已经形成的知识关联。对于改写询问、补全诱导、上下文诱导等攻击形式，外部过滤策略可能出现漏检。模型编辑则尝试在参数层面对特定事实关联进行局部修改，为隐私知识清洗提供了不同于重训和输出过滤的技术路径。

## 1.2 大语言模型隐私泄露问题

隐私泄露并不只表现为模型在直接询问时输出敏感值。用户可以通过改写问题、要求模型补全个人资料、设定角色扮演场景或提供上下文诱导模型继续生成。电话号码、邮箱等结构化隐私值具有明显格式特征，既便于自动检测，也可能在续写式提示中被模型生成。因而，隐私清洗评价需要覆盖多种攻击类型，而不能只依赖单一问答模板。

真实个人信息存在合规和复现困难：研究者很难确认某条真实隐私是否确实被预训练模型记忆，也不应在论文中展示真实个人敏感信息。本文采用合成隐私事实，目的不是模拟全部真实世界复杂性，而是构造可控、可归因、可重复的隐私泄露对象，使不同编辑方法能够在同一实验起点上比较。

## 1.3 模型编辑用于隐私清洗的动机

模型编辑旨在以较低代价修改模型中的特定知识关联。与全量重训相比，编辑方法通常只处理少量目标事实，计算成本更低；与输出过滤相比，编辑方法直接作用于模型参数或模型调用路径，更接近隐私知识清洗的目标。普通事实编辑通常希望把错误事实替换为正确事实，而隐私清洗希望模型对目标隐私查询不输出敏感值，同时保留相关公开事实，这使隐私清洗天然具有多目标属性。

## 1.4 当前方法的不足

现有模型编辑评测通常强调可靠性、泛化性和局部性，但隐私拒答式清洗还需要显式衡量公开问题是否被误拒。若只统计隐私查询上的拒答比例，模型可能通过对整类人物问题泛化拒答获得较好数值；若只报告公开知识保持，又可能忽略目标隐私泄露。因此，本文同时报告隐私值包含率、PII 格式匹配率、公开知识保持率和公开问题拒答率，用于区分真正隐私压制和过度拒答。

## 1.5 本文研究内容与贡献

本文围绕“基于模型编辑的大语言模型隐私知识清洗”构建完整实验流程。主要贡献如下。

第一，构建 private/public 解耦的合成隐私清洗基准，支持同时评估目标隐私压制与公开知识保持。该基准包含隐私事实、同主体公开事实、同关系其他主体公开事实和通用知识，使副作用分析具有明确对象。

第二，基于 LoRA 构造可控隐私泄露模型，并将 ROME、MEMIT、FT 等模型编辑方法适配到隐私拒答式清洗任务。该设计避免真实预训练 PII 来源不可控的问题，使编辑效果能够追溯到明确的合成事实集合。

第三，设计 PACE、CAPE 与 CAPE-Anchor 闭环请求选择策略，用于分析 residual leakage re-edit、过度拒答和公开知识锚点约束。三者均位于编辑器外层，不改变 ROME/MEMIT 的底层更新公式。

第四，建立隐私压制—知识保持权衡分析，包括 Public Refusal、attack-type split、Qwen public transfer 和 GPT-J boundary check。该评估体系将正向结果与边界结果同时纳入论文叙事，避免把单指标最优误读为全面成功。

## 1.6 论文组织结构

第2章介绍大语言模型隐私泄露、模型编辑和隐私清洗评价基础；第3章给出合成隐私基准、LoRA 注入和 PACE/CAPE/CAPE-Anchor 方法设计；第4章报告实验设置、主结果、消融、公开迁移与边界分析；第5章说明系统实现与软件组织；第6章总结全文并讨论局限与后续工作。

# 第2章 相关技术与研究基础

## 2.1 大语言模型与参数化知识

大语言模型可被视为在参数中编码了大量语言模式和事实关联。部分研究指出 Transformer 前馈层具有类似键值记忆的作用，注意力层和前馈层共同参与知识存储与调用。隐私记忆可以理解为模型在参数中形成的关于个人敏感事实的可触发关联。当查询提示与该关联匹配时，模型可能输出对应敏感值。

参数化知识并不以显式数据库记录形式存在，而是分布在大量权重和激活模式中。隐私事实一旦进入这种表示空间，删除原始训练样本并不必然消除模型输出该事实的能力。模型编辑的研究价值正来自这一点：它尝试在较低重训成本下修改局部知识关联，从而为后训练隐私清洗提供可操作路径。

## 2.2 隐私泄露、训练数据记忆与拒答机制

训练数据抽取和记忆研究表明，神经语言模型可能在特定条件下复现训练样本中的长片段或结构化信息[1-3]。差分隐私和机器遗忘从训练机制或样本移除角度控制隐私风险[4-5]，但往往需要重新训练、访问原始训练数据或付出较高效用代价。对于已经部署或已经后训练的模型，参数级局部编辑是一条更轻量的补救路径。

拒答机制本身并不等价于隐私清洗。合理拒答应发生在目标隐私查询上，而公开知识、非敏感事实和通用问题不应被误拒。因此，本文把公开问题拒答率作为核心副作用指标，用于判断模型是否只是学到宽泛拒答模板。

## 2.3 模型编辑方法概述

模型编辑旨在不全量重训模型的情况下修改特定知识。ROME、MEMIT、MEND、SERAC、IKE、GRACE 等方法分别从局部参数更新、批量编辑、元学习、外部记忆、上下文学习或适配器角度探索编辑机制[7-11]。EasyEdit 将多种编辑方法统一到同一工程框架中，为同模型、同数据、同评估协议下的横向比较提供了基础[14-15]。

## 2.4 ROME 与 MEMIT

ROME 面向单条事实编辑，通过定位事实关联在模型中的关键层并执行秩一更新来改变目标输出[7]。MEMIT 面向批量事实编辑，在多个层上估计更新并同时注入多条编辑约束[8]。二者的原始设计目标并非隐私清洗，因此需要结合公开知识保持和过度拒答指标重新评价。

## 2.5 微调类与定位类编辑基线

FT 通过小规模梯度更新直接优化目标请求，能够提供强编辑端点，但可能导致过覆盖。KN 等定位类方法尝试寻找知识相关神经元，但在大模型和长上下文下计算开销较高。本文保留 FT 作为强覆盖对照，并将 KN/IKE 的资源受限结果纳入 failure/resource limitation 表，而不伪造成有效对比。

## 2.6 隐私清洗评价指标

本文主要使用六类指标。Private Value Contains 判断目标隐私值是否出现在私有攻击输出中；PII-format Regex 统计电话、邮箱等结构化敏感格式；Sensitive Pattern 统计敏感格式输出倾向；Private Refusal 衡量隐私查询拒答；Public Contains 衡量公开事实是否被保留；Public Refusal 衡量公开问题是否被误拒。公开 CounterFact/zsRE 实验只报告 Reliability、Generalization 和 Locality，不与 synthetic privacy 指标混表。

## 2.7 本章小结

本章说明了隐私知识清洗不同于普通事实编辑：其目标不是替换一个事实，而是在压制敏感值的同时保留公开知识。该任务需要同时分析泄露、保持和过度拒答，为后续方法设计提供了评价基础。

# 第3章 方法设计

## 3.1 任务定义

设原模型为 M，编辑后模型为 M'，待清洗隐私事实集合为 K_F，需要保留的公开事实集合为 K_R。隐私知识清洗要求 M' 在隐私查询集合 Q_F 上降低目标敏感值输出，在公开查询集合 Q_R 上保持正确公开事实，并避免把非敏感问题泛化为拒答。本文使用如下多目标形式描述任务：

$$\\mathcal{{J}}(M') = \\operatorname{{Leakage}}(M', K_F) + \\lambda \\operatorname{{UtilityLoss}}(M, M', K_R) + \\mu \\operatorname{{OverRefusal}}(M', K_R).$$

该公式不是训练目标，而是评价视角：隐私泄露越低越好，公开知识损失越小越好，公开问题过度拒答越少越好。

## 3.2 private/public 解耦的合成隐私基准

评测基准包含 100 个虚拟人物。每个人物具有电话号码、邮箱等隐私事实，以及职业、学校、雇主、家乡等公开事实。公开事实进一步划分为同主体公开知识、同关系其他主体公开知识和通用知识，用于观察编辑是否只影响目标隐私，还是扩散到相邻公开知识。

表1给出数据构成。合成数据的优势在于 ground truth 完全可控，能够精确判断模型输出是否包含目标隐私值或公开答案。与真实 PII 相比，合成数据避免了展示真实个人信息的伦理和合规问题，同时保留电话号码、邮箱等结构化隐私格式，使泄露检测具有可重复性。

表1 合成隐私基准数据构成

|项目|规模或说明|
|---|---|
|人物数量|100|
|隐私事实类型|电话、邮箱|
|公开事实类型|职业、学校、雇主、家乡等|
|隐私攻击提示|direct、paraphrase、completion、roleplay、context|
|完整隐私评估|3000 条生成|
|公开知识评估|2520 条生成|

## 3.3 LoRA 可控隐私注入

本文使用 MLP-only LoRA 将隐私事实和公开事实共同注入 Qwen2.5-7B，并将适配器合并为待编辑模型。合并后模型在隐私查询上具有较高泄露率，同时在公开问题上保持较高包含式保持率，说明该模型适合作为统一编辑起点。

选择合并后的模型作为编辑对象还有一个重要原因：ROME 和 MEMIT 在运行时直接修改模型参数，如果继续依赖外部 adapter，会增加编辑对象和评估路径的不确定性。将 LoRA adapter 合并后，所有方法面对的是同一个普通 causal language model，便于比较编辑前后的参数级影响。

## 3.4 隐私拒答式模型编辑

第一轮编辑将隐私事实对应请求改写为拒答目标，使模型在目标隐私查询上倾向于拒绝输出敏感值。对于电话号码或邮箱，目标不是把旧值替换为另一个值，而是让模型避免输出该敏感值。因此编辑请求将原始隐私查询映射到安全拒答表达。该设计能够直接评估编辑器是否可以把事实编辑机制迁移到隐私保护目标。

## 3.5 PACE：隐私感知闭环编辑

PACE 是仅基于残余泄露风险的闭环再编辑策略。其核心思想是在第一轮编辑后重新评估残余泄露，从仍然泄露的样本中构造第二轮编辑请求，再执行补充编辑。PACE 保持底层编辑器公式不变，通过残余泄露驱动请求构造。本文采用二轮实例化，是为了控制运行成本并观察过度拒答风险。

## 3.6 CAPE：副作用感知请求选择

CAPE 是引入局部性风险、主体预算和请求数量控制的副作用感知请求筛选策略。它从残余泄露候选中选择请求，同时限制每个主体和总请求数，避免把拒答模式扩散到保留集合。其请求评分可概括为：

$$\\operatorname{{score}}(q)=\\operatorname{{Risk}}_{{leak}}(q)-\\lambda \\operatorname{{Risk}}_{{loc}}(q).$$

## 3.7 CAPE-Anchor：公开知识锚点约束

CAPE-Anchor 在 CAPE 的基础上显式加入公开知识锚点。它不修改 ROME/MEMIT/EasyEdit 底层算法，而是在二轮请求选择阶段将 public anchors 纳入最终请求集合：

$$R_{{final}}=R_{{round1}}\\cup R_{{privacy}}\\cup R_{{anchor}}.$$

其中 R_privacy 来自 residual leakage，R_anchor 来自需要保留的公开事实。该设计的目的不是追求最低隐私泄露率，而是检验公开锚点能否将 PACE/CAPE 的公开知识塌缩拉回到更可用的折中区域。

## 3.8 评价指标与权衡分析

本文将隐私值包含率作为目标隐私泄露的主要自动指标。公开知识包含式保持率衡量公开事实答案是否出现在公开问题输出中。公开问题拒答率统计公开问题中出现拒答标记的比例，用于衡量过度拒答副作用。三类指标之间存在天然张力，因此本文不使用单一排名，而使用隐私压制—知识保持二维图和 failure/resource 表共同解释结果。

## 3.9 本章小结

本章给出了从合成数据、LoRA 注入、隐私拒答编辑到 PACE/CAPE/CAPE-Anchor 请求选择的完整方法设计。核心思想是把隐私清洗从单纯拒答任务转化为带公开知识约束的多目标编辑问题。

# 第4章 实验设计与结果分析

## 4.1 实验设置

主实验使用 Qwen privacy merged 模型和 synthetic privacy v2 数据集。所有方法在同一隐私事实、公开事实和攻击提示集合上评估。公开 CounterFact/zsRE 实验只用于 factual editing 迁移验证，不作为 PII 清洗主结果。GPT-J-6B 仅作为第二模型边界验证，不写成跨模型成功。

## 4.2 数据集与模型

表2给出实验平台与运行环境。模型权重、缓存和大规模生成文件不作为论文正文的一部分，论文只保留小型 summary、CSV、图表和审计结果。

表2 实验平台与运行环境

|项目|说明|
|---|---|
|主实验模型|Qwen2.5-7B privacy merged|
|公开迁移模型|Qwen2.5-7B|
|边界验证模型|GPT-J-6B|
|主数据集|synthetic privacy v2|
|公开数据集|CounterFact、zsRE|
|主要框架|EasyEdit、Transformers、PyTorch|
|主要硬件|48GB GPU 实例|
|结果状态|2026-06-23 frozen artifacts|

## 4.3 评价指标

Synthetic privacy 主表报告隐私值包含率、PII 格式匹配率、敏感格式输出率、隐私拒答率和公开知识保持率。公开数据集报告 Reliability、Generalization 和 Locality。所有 failed/resource-limited 项单独记录，不将 failed 写成 0，也不将 missing 留空解释为有效结果。

图1展示本文整体实验路线。该流程从合成数据构造和 LoRA 注入开始，经标准编辑器与闭环请求选择，最终输出隐私压制、公开知识保持和边界分析结果。

图1 总体技术路线图

![图1 总体技术路线图](artifacts/final_paper_assets_20260623/figures/fig1_framework_pipeline.png)

图2展示 PACE、CAPE 与 CAPE-Anchor 的请求构造差异。PACE 只关注残余泄露，CAPE 增加副作用感知筛选，CAPE-Anchor 进一步显式加入公开知识锚点。

图2 PACE/CAPE/CAPE-Anchor 请求构造机制图

![图2 PACE/CAPE/CAPE-Anchor 请求构造机制图](artifacts/final_paper_assets_20260623/figures/fig2_pace_cape_anchor_mechanism.png)

## 4.4 Synthetic privacy 主实验

表3给出 Qwen synthetic privacy 主结果。合并后泄露模型的隐私值包含率为 0.9387，公开知识保持率为 0.9766，说明模型起点同时具有可测隐私泄露和较稳定公开知识。提示拒答的隐私值包含率仍为 0.9490，说明单纯提示层拒答基本不能改变模型内部记忆。ROME 将隐私值包含率降至 0.5787，但公开知识保持率也下降到 0.5591；MEMIT 保持公开知识较好，公开知识保持率为 0.8472，但隐私值包含率仍为 0.8140。FT 将隐私值包含率压到 0.0000，但公开知识保持率仅为 0.0040，表现为强覆盖端点。

表3 Synthetic privacy 主结果（↓越低越好，↑越高越好）

{md_table(synthetic, ['方法', '状态', '隐私值包含率↓', 'PII格式率↓', '隐私拒答率', '公开知识保持率↑'])}

PACE 和 CAPE 分别将隐私值包含率压低到 0.0243 和 0.0443，但公开知识保持率仅为 0.0984 和 0.1119。这说明残余泄露驱动的闭环再编辑能够显著压制隐私输出，但也会造成 public collapse。图3将各方法放入隐私压制—知识保持二维空间：左上角代表强隐私压制但低公开保持，右上角才是理想区域。当前结果显示没有单一方法同时达到两个目标的最优区域，方法差异更适合解释为 trade-off。

图3 Synthetic privacy 隐私压制—知识保持权衡散点图

![图3 Synthetic privacy 隐私压制—知识保持权衡散点图](artifacts/final_paper_assets_20260623/figures/fig3_synthetic_privacy_utility_tradeoff.png)

## 4.5 PACE/CAPE 对比分析

PACE 的意义在于证明 residual leakage re-edit 可以显著压低隐私泄露，但其副作用同样明显。CAPE 在 PACE 基础上加入副作用感知请求筛选，能够相对缓解极端过度拒答，但仍未完全恢复公开知识保持。该结果说明，隐私清洗不能只追求隐私值包含率最低；若公开事实也被拒答，系统可用性会显著下降。

从论文主张看，PACE/CAPE 不是“全面优于”标准编辑器，而是揭示了闭环请求扩展的两面性：它能补足单轮编辑覆盖不足，也会放大拒答目标的扩散风险。这一发现直接引出 CAPE-Anchor。

## 4.6 CAPE-Anchor 消融实验

表4给出 CAPE-Anchor 消融结果。K0 是无公开锚点的受限再编辑对照；K1/K2 引入公开锚点后，公开知识保持率明显上升。PACE-Lite K0 的公开知识保持率为 0.4210，CAPE-Anchor K1 提升到 0.6008，K2 进一步提升到 0.6833。与此同时，隐私值包含率也从 0.3323 上升到 0.4603 和 0.6357，说明公开锚点带来更好的 utility，但会牺牲部分 privacy。

表4 CAPE-Anchor 消融结果

{md_table(cape, ['配置', '隐私值包含率↓', 'PII格式率↓', '隐私拒答率', '公开知识保持率↑', '解释'])}

图4展示了 CAPE-Anchor 的折中变化。K1 通常是较合理的折中点：它比 PACE/CAPE 明显提高公开知识保持，同时隐私值包含率仍低于 ROME。K2 的公开保持更高，但隐私压制更弱。因此，本文采用 Claim A 的限定写法：CAPE-Anchor 形成有限有效改进，即相对于 naive PACE/CAPE 提供更合理的隐私压制—知识保持折中，而不是完全解决隐私清洗。

图4 CAPE-Anchor 消融折中图

![图4 CAPE-Anchor 消融折中图](artifacts/final_paper_assets_20260623/figures/fig4_cape_anchor_ablation_tradeoff.png)

## 4.7 Qwen 公开事实编辑迁移验证

CounterFact 和 zsRE 不是 PII 清洗任务。本文将其作为公开事实编辑迁移验证，用于观察 closed-loop request selection 在非隐私 factual editing 场景下的外部行为。表5显示，Qwen 上 ROME/FT 是较强 baseline，wrapper 可以迁移运行但效果中等。CounterFact 上 ROME+PACE-Edit 和 ROME+CAPE-Edit 的可靠性均为 0.5179；zsRE 上二者可靠性均为 0.3117。这说明 wrapper 不是公开事实编辑上的强正结果，只能作为迁移验证和边界说明。

表5 Qwen public transfer 结果

{md_table(qwen, ['数据集', '方法', '样本数', '可靠性↑', '泛化性↑'])}

图5给出 Qwen public transfer 的可视化结果。该图应被解释为公开事实编辑压力测试，而不是隐私清洗成功证明。公开数据集结果与 synthetic privacy 主实验回答的是不同问题，不能混成同一排行榜。

图5 Qwen public transfer 对比图

![图5 Qwen public transfer 对比图](artifacts/final_paper_assets_20260623/figures/fig5_public_transfer_qwen.png)

## 4.8 GPT-J 第二模型边界验证

GPT-J-6B 上的第二模型验证表明，ROME 与 FT baseline 在 CounterFact 和 zsRE 上均能正常运行，但未重新调参的 ROME-based PACE/CAPE wrapper 几乎完全失效。CounterFact 上 wrapper 的 rewrite 成功数为 0/220，zsRE 上为 1/216。该现象已经通过 per-case 文件复核，并非路径混用或汇总脚本误读，说明闭环请求扩展策略在跨模型迁移时对底层编辑器超参数和请求规模较为敏感。

表6 GPT-J second-model boundary check

{md_table(gptj, ['数据集', '方法', '可靠性↑', '泛化性↑', '逐样本rewrite', '解释'])}

图6直观展示了 GPT-J 的边界现象：ROME/FT baseline 保持较高可靠性，而未调参 wrapper 接近 0。这一结果不能写成跨模型成功，但对论文有价值，因为它明确了方法边界，并提示后续需要 model-specific calibration、locality-aware constraint 或 retain-aware objective。

图6 GPT-J wrapper boundary check 图

![图6 GPT-J wrapper boundary check 图](artifacts/final_paper_assets_20260623/figures/fig6_gptj_wrapper_boundary_check.png)

## 4.9 失败项与资源限制分析

表7列出失败项和资源限制。KN 在部分设置中因显存压力失败或结果极弱；IKE 因缺少 SentenceTransformer 依赖未形成有效结果；GPT-J wrapper 虽然运行完成，但作为负向边界结果处理。这些内容不应被隐藏，也不应写成工程事故，而应作为资源边界和方法边界说明。

表7 Failure and resource limitation

{md_table(failures, ['范围', '方法/现象', '状态', '原因', '论文处理'])}

图7总结了主要失败和边界项。该图的作用是保证论文结果口径透明：failed/missing 不伪造成 0，negative result 不伪造成 success。

图7 Failure/resource limitation summary

![图7 Failure/resource limitation summary](artifacts/final_paper_assets_20260623/figures/fig7_failure_matrix_summary.png)

## 4.10 案例分析

典型案例可以从三个现象理解本文结果。第一，提示拒答并不能可靠消除模型内部记忆，模型仍可能在补全或改写提示下输出敏感值。第二，PACE/CAPE 的强隐私压制往往伴随公开问题拒答扩散，说明残余泄露驱动再编辑不能脱离公开知识约束。第三，CAPE-Anchor 通过加入公开锚点，把方法从强拒答端点拉回到更可用的折中区域，但公开保持提高并不意味着隐私压制没有代价。

对于 GPT-J，案例层面的关键信息是：相同模型上的 ROME/FT baseline 正常，只有未调参 wrapper 塌缩。因此该结果更适合放在边界分析中，而不是作为否定主实验价值的证据。

## 4.11 本章小结

实验结果支持本文的限定主张：CAPE-Anchor 在 Qwen synthetic privacy 主任务上形成有限有效改进，表现为更合理的隐私压制—知识保持折中。公开 Qwen 实验提供中性迁移验证，GPT-J 实验提供负向边界证据。整体结论不是“某一方法全面最优”，而是隐私清洗必须同时分析泄露、公开保持和过度拒答。

# 第5章 系统设计与软件实现

## 5.1 系统总体结构

系统包含数据构造、LoRA 注入、编辑请求构造、模型编辑运行、隐私生成评估、公开知识评估和结果汇总模块。各模块通过 JSON、JSONL、CSV 和 Markdown artifact 连接，便于复查输入输出。

## 5.2 数据构造模块

数据构造模块生成虚拟人物、隐私事实、公开事实和多种攻击提示。每条样本保留 case_id、attack_type、template_id 和目标值，确保后续生成结果能够追溯到具体人物和提示类型。

## 5.3 隐私注入模块

隐私注入模块通过 LoRA 训练数据构造和 adapter merge 形成统一待编辑模型。该模块只负责构造可控泄露对象，不参与后续编辑效果评价，避免把注入效果和清洗效果混淆。

## 5.4 模型编辑模块

模型编辑模块基于 EasyEdit 调用 ROME、MEMIT、FT 等方法，并通过 PACE/CAPE/CAPE-Anchor 构造不同请求集合。底层编辑器与外层请求选择分离，使实验能够回答“编辑器能力”和“请求选择副作用”两个问题。

## 5.5 评估与结果分析模块

评估模块对 private attacks 和 public retain prompts 分别生成并统计指标。公开 CounterFact/zsRE 结果由独立脚本汇总，避免 public factual editing 指标和 synthetic privacy 指标混表。最终论文资产由冻结 CSV 生成，确保每个表格和图都能追溯到具体 artifact。

## 5.6 实验平台与运行环境

大模型实验在 Linux GPU 环境中运行，本地 Windows 环境负责代码维护、结果整理和论文生成。模型权重、HF 缓存和大规模生成中间文件不纳入最终论文仓库；仓库保留脚本、小型 summary、CSV、图表和文档。

## 5.7 代码组织与软件使用说明

主要脚本分布在 scripts 目录中，包括 synthetic privacy 数据生成、LoRA 训练数据构造、隐私拒答编辑、PACE/CAPE 请求构造、公开 benchmark 运行、结果抽取和最终论文资产生成。实验完成后，最终表格位于 artifacts/final_paper_assets_20260623/tables，最终图位于 artifacts/final_paper_assets_20260623/figures。

## 5.8 本章小结

系统实现强调可审计性和可复现性。通过保存请求、评估、summary、failure matrix 和最终表图，论文中的主要结论可以在不重新运行 GPU 实验的情况下复核。

# 第6章 总结与展望

## 6.1 工作总结

本文围绕基于模型编辑的大语言模型隐私知识清洗问题，构建了从合成隐私评测基准、LoRA 可控隐私注入、模型编辑执行到副作用审计的完整实验闭环。该流程使待清洗隐私事实、需要保留的公开事实以及多类攻击提示均具有明确 ground truth，从而能够在统一协议下比较不同编辑策略。

## 6.2 主要结论

实验表明，不同编辑策略在隐私压制和知识保持两个目标上呈现差异化端点。Prompt Refusal 基本不能改变模型内部记忆；MEMIT 表现出较强公开知识保持能力，但压制强度有限；ROME 形成中间区域；FT、PACE 和 CAPE 构成强隐私压制但公开保持较弱的端点；CAPE-Anchor 通过公开锚点把 naive PACE/CAPE 的 public collapse 拉回到更合理区域。公开 Qwen 结果说明 wrapper 可以迁移运行但效果中等，GPT-J 结果揭示跨模型迁移边界。

## 6.3 局限性分析

本文的合成隐私基准提供可控 ground truth，但不能覆盖真实世界 PII 的全部语义复杂性。自动 contains 和 regex 指标适合稳定比较，但不能替代人工语义评估。CAPE-Anchor 当前仍是请求层约束，尚未在底层编辑目标中显式加入 retain loss。GPT-J wrapper collapse 也表明，闭环请求扩展对模型结构、ROME 超参数和请求集合规模敏感。

## 6.4 后续工作

后续研究可以从三个方向推进。第一，引入模型特定校准和 locality-aware constraint，使闭环 wrapper 在不同模型上更稳定。第二，在编辑目标中加入 retain-aware objective，把公开知识锚点从请求构造扩展为显式优化约束。第三，结合人工评估和更多真实场景数据，对拒答合理性、公开回答质量和多轮攻击鲁棒性进行更细粒度分析。

# 参考文献

[1] Carlini N., Tramer F., Wallace E., et al. Extracting Training Data from Large Language Models. USENIX Security Symposium, 2021.

[2] Carlini N., Ippolito D., Jagielski M., et al. Quantifying Memorization Across Neural Language Models. ICLR, 2023.

[3] Kandpal N., Wallace E., Raffel C. Deduplicating Training Data Mitigates Privacy Risks in Language Models. ICML, 2022.

[4] Dwork C., Roth A. The Algorithmic Foundations of Differential Privacy. Foundations and Trends in Theoretical Computer Science, 2014.

[5] Bourtoule L., Chandrasekaran V., Choquette-Choo C. A., et al. Machine Unlearning. IEEE Symposium on Security and Privacy, 2021.

[6] Hu E. J., Shen Y., Wallis P., et al. LoRA: Low-Rank Adaptation of Large Language Models. ICLR, 2022.

[7] Meng K., Bau D., Andonian A., Belinkov Y. Locating and Editing Factual Associations in GPT. NeurIPS, 2022.

[8] Meng K., Sharma A. S., Andonian A., Belinkov Y., Bau D. Mass-Editing Memory in a Transformer. ICLR, 2023.

[9] Mitchell E., Lin C., Bosselut A., Manning C. D., Finn C. Fast Model Editing at Scale. ICLR, 2022.

[10] Mitchell E., Lin C., Bosselut A., Finn C., Manning C. D. Memory-Based Model Editing at Scale. ICML, 2022.

[11] De Cao N., Aziz W., Titov I. Editing Factual Knowledge in Language Models. EMNLP, 2021.

[12] Dai D., Dong L., Hao Y., Sui Z., Chang B., Wei F. Knowledge Neurons in Pretrained Transformers. ACL, 2022.

[13] Geva M., Schuster R., Berant J., Levy O. Transformer Feed-Forward Layers Are Key-Value Memories. EMNLP, 2021.

[14] Wang P., Zhang N., Tian B., et al. EasyEdit: An Easy-to-use Knowledge Editing Framework for Large Language Models. ACL Demo, 2024.

[15] Yao Y., Wang P., Tian B., et al. Editing Large Language Models: Problems, Methods, and Opportunities. EMNLP Findings, 2023.

[16] Petroni F., Rocktaschel T., Riedel S., et al. Language Models as Knowledge Bases? EMNLP-IJCNLP, 2019.

[17] Levy O., Seo M., Choi E., Zettlemoyer L. Zero-Shot Relation Extraction via Reading Comprehension. CoNLL, 2017.

[18] Biderman S., Schoelkopf H., Anthony Q., et al. Pythia: A Suite for Analyzing Large Language Models Across Training and Scaling. ICML, 2023.

[19] Gao L., Biderman S., Black S., et al. The Pile: An 800GB Dataset of Diverse Text for Language Modeling. arXiv:2101.00027, 2020.

[20] Klimt B., Yang Y. The Enron Corpus: A New Dataset for Email Classification Research. ECML, 2004.

[21] Eldan R., Russinovich M. Who's Harry Potter? Approximate Unlearning in LLMs. arXiv:2310.02238, 2023.

[22] Maini P., Feng Z., Schwarzschild A., et al. TOFU: A Task of Fictitious Unlearning for LLMs. arXiv:2401.06121, 2024.

[23] Wu X., Li J., Xu M., et al. DEPN: Detecting and Editing Privacy Neurons in Pretrained Language Models. EMNLP, 2023.

# 附录

附录A列出最终结果资产位置：`artifacts/final_paper_assets_20260623/`。附录B列出 GPT-J per-case 审计文件：`artifacts/final_comparison_20260623_complete/gptj_per_case_audit.csv`。附录C列出 failure/resource limitation 表，用于说明未扩展 GPU 实验的原因。
"""


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if len(str(text)) < 18 else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(str(text))
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(8.5)
    run.bold = bold
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_table(doc: Document, rows: List[List[str]]) -> None:
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j, value in enumerate(row):
            cell = table.cell(i, j)
            set_cell_text(cell, value, bold=(i == 0))
            if i == 0:
                set_cell_shading(cell, "E8EEF5")
    doc.add_paragraph()


def add_image(doc: Document, path: str, width_cm: float = 14.5) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(Path(path)), width=Cm(width_cm))


def configure_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    styles = doc.styles
    styles["Normal"].font.name = "宋体"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    styles["Normal"].font.size = Pt(10.5)
    for style_name, size, color in [
        ("Heading 1", 16, RGBColor(46, 116, 181)),
        ("Heading 2", 13, RGBColor(46, 116, 181)),
        ("Heading 3", 12, RGBColor(31, 77, 120)),
    ]:
        style = styles[style_name]
        style.font.name = "黑体"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        style.font.size = Pt(size)
        style.font.color.rgb = color


def parse_md_table(lines: List[str], start: int) -> tuple[List[List[str]], int]:
    rows = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        line = lines[i].strip()
        if not re.match(r"^\|[-:| ]+\|$", line):
            cells = [c.strip() for c in line.strip("|").split("|")]
            rows.append(cells)
        i += 1
    return rows, i


def markdown_to_docx(md: str, out_path: Path) -> None:
    doc = Document()
    configure_doc(doc)
    lines = md.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        if not line:
            i += 1
            continue
        if line.startswith("# "):
            text = line[2:].strip()
            if text.startswith("第") or text in {"摘要", "关键词", "参考文献", "附录"}:
                doc.add_heading(text, level=1)
            else:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(text)
                run.bold = True
                run.font.name = "黑体"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
                run.font.size = Pt(18)
            i += 1
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
            i += 1
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
            i += 1
            continue
        if line.startswith("!["):
            match = re.match(r"!\[(.*?)\]\((.*?)\)", line)
            if match:
                add_image(doc, match.group(2))
            i += 1
            continue
        if line.startswith("|"):
            rows, i = parse_md_table(lines, i)
            if rows:
                add_table(doc, rows)
            continue
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Pt(21)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(line)
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        run.font.size = Pt(10.5)
        i += 1
    doc.save(out_path)


def write_changelog() -> None:
    CHANGELOG.write_text(
        """# Course Paper V7 Final Changelog

Last updated: 2026-06-23
Final result status: frozen
No further GPU expansion unless explicitly approved

## 主要变更

1. 以 v5 完整论文为骨架，删除 v6 结果占位、待填和过期实验计划。
2. 将最终 claim 固定为 Claim A：CAPE-Anchor 形成有限有效改进，但仅限 Qwen synthetic privacy 主实验。
3. 用 frozen artifacts 更新 Synthetic privacy、CAPE-Anchor、Qwen public transfer、GPT-J boundary 和 failure/resource 表。
4. 将 GPT-J 从“跨模型成功”改为“second-model boundary check”，明确 wrapper collapse 已经 per-case 审计。
5. 整合最终图表资产 `artifacts/final_paper_assets_20260623/figures`，并为每张图加入解释段。
6. 生成 `课程论文_v7_final_academic_polished.docx` 与 `docs/course_paper_v7_final_academic_polished.md`，不覆盖旧版本。

## 删除或降级内容

- 删除 TOFU、Enron、LLaMA、Pythia 等未完成扩展作为“下一步必做”的表述。
- 删除“结果待填”“Claim A/B/C 待选择”等过期口径。
- 不再把 public benchmark 写成 PII 清洗证明。
- 不再把 GPT-J wrapper 失败写成工程事故或主方法失败。
""",
        encoding="utf-8",
    )


def write_checklist() -> None:
    CHECKLIST.write_text(
        """# Paper Final Quality Checklist

Last updated: 2026-06-23
Final result status: frozen
No further GPU expansion unless explicitly approved

## 结果口径

- [x] 主实验限定为 Qwen privacy merged × synthetic privacy v2。
- [x] Claim A 写成“有限有效改进 / 更合理 privacy-utility trade-off”。
- [x] 未声称 CAPE-Anchor 隐私泄露最低、无损保持或全面优于 baseline。
- [x] Qwen public transfer 写成公开事实编辑迁移验证。
- [x] GPT-J 写成 second-model boundary check，不写成跨模型成功。
- [x] failed/missing 未写成 0。

## 图表

- [x] 图1 总体技术路线图。
- [x] 图2 PACE/CAPE/CAPE-Anchor 请求构造机制图。
- [x] 图3 Synthetic privacy 隐私压制—知识保持权衡图。
- [x] 图4 CAPE-Anchor 消融折中图。
- [x] 图5 Qwen public transfer 对比图。
- [x] 图6 GPT-J wrapper boundary check 图。
- [x] 图7 Failure/resource limitation summary。
- [x] 表格包含状态或失败原因，不隐藏资源限制。

## 写作

- [x] 摘要避免工程过程词和过强主张。
- [x] 贡献点限定为数据、方法适配、闭环请求选择和评估体系。
- [x] 每个主要结果段绑定具体指标。
- [x] 结论包含局限性和后续工作，但不写成自我否定。

## 待人工最终检查

- [ ] Word 中目录、页码和学校模板格式是否需要手动套用。
- [ ] 参考文献会议/期刊信息是否需要按老师要求进一步核验。
- [ ] DOCX 中图片尺寸是否符合最终提交模板。
- [ ] 本机未检测到 LibreOffice/soffice，DOCX 已完成结构审计，但仍建议用 Word 打开做最终视觉检查。
""",
        encoding="utf-8",
    )


def main() -> int:
    tables = load_tables()
    md = build_markdown(tables)
    OUT_MD.write_text(md, encoding="utf-8")
    markdown_to_docx(md, OUT_DOCX)
    write_changelog()
    write_checklist()
    print(f"wrote {OUT_MD}")
    print(f"wrote {OUT_DOCX}")
    print(f"wrote {CHANGELOG}")
    print(f"wrote {CHECKLIST}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
